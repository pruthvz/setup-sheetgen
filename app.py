"""
SheetGen
Copyright (c) 2024–2025 Pruthvi. All rights reserved.
Unauthorized copying, distribution, or modification is prohibited.
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext, simpledialog
import threading
import json
import os
import re
import sys
from pathlib import Path
from datetime import date

# Suppress tqdm progress bar (used by docx2pdf) — avoids terminal/console spam when building exe
os.environ["TQDM_DISABLE"] = "1"

# PyInstaller windowed exe: stderr is None; docx2pdf/tqdm crashes without this
if sys.stderr is None:
    sys.stderr = open(os.devnull, "w")

# ── optional heavy deps ───────────────────────────────────────────────────
try:
    import openpyxl
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    Pt = RGBColor = None

try:
    from docx2pdf import convert as docx2pdf_convert
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

# Optional: OCR for scanned PDFs (Rename tool)
try:
    import fitz  # pymupdf
    from PIL import Image
    import io
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import pytesseract
    HAS_PYTESSERACT = True
except ImportError:
    HAS_PYTESSERACT = False


def _configure_tesseract_for_exe():
    """
    Point pytesseract to Tesseract so OCR works without user installing it globally.
    - When running as PyInstaller exe: use bundled tesseract from _MEIPASS
    - When running from Python: use tesseract_bundle/ in project folder (must include all DLLs)
    - Fallback: use system Tesseract if installed in Program Files
    """
    if not HAS_PYTESSERACT:
        return
    try:
        base = getattr(sys, "_MEIPASS", None)
        if base:
            # Running as exe: use bundled tesseract
            tesseract_dir = os.path.join(base, "tesseract")
            exe_path = os.path.join(tesseract_dir, "tesseract.exe")
            tessdata_dir = os.path.join(tesseract_dir, "tessdata")
        else:
            # Running from Python: use tesseract_bundle in project folder
            app_dir = os.path.dirname(os.path.abspath(__file__))
            tesseract_dir = os.path.join(app_dir, "tesseract_bundle")
            exe_path = os.path.join(tesseract_dir, "tesseract.exe")
            tessdata_dir = os.path.join(tesseract_dir, "tessdata")
        if os.path.isfile(exe_path) and os.path.isdir(tessdata_dir):
            pytesseract.pytesseract.tesseract_cmd = exe_path
            # Tesseract 5.x expects TESSDATA_PREFIX = path to tessdata folder (or parent + /)
            os.environ["TESSDATA_PREFIX"] = tessdata_dir + os.sep
            return
        # Fallback: system Tesseract (Program Files)
        for pf in (os.environ.get("ProgramFiles", "C:\\Program Files"),
                   os.environ.get("ProgramFiles(x86)", "C:\\Program Files (x86)")):
            sys_path = os.path.join(pf, "Tesseract-OCR", "tesseract.exe")
            if os.path.isfile(sys_path):
                pytesseract.pytesseract.tesseract_cmd = sys_path
                os.environ["TESSDATA_PREFIX"] = os.path.join(pf, "Tesseract-OCR")
                return
    except Exception:
        pass


_configure_tesseract_for_exe()

try:
    import anthropic as _anthropic_lib
    HAS_ANTHROPIC = True
except ImportError:
    HAS_ANTHROPIC = False

try:
    import openai as _openai_lib
    HAS_OPENAI = True
except ImportError:
    HAS_OPENAI = False


# ── colours ───────────────────────────────────────────────────────────────
BG      = "#0d0d12"
PANEL   = "#16161f"
PANEL2  = "#1e1e2a"
ACCENT  = "#00e5ff"
GREEN   = "#00ff9d"
PURPLE  = "#c084fc"
TEXT    = "#e8e8f0"
MUTED   = "#55556a"
BORDER  = "#252535"
YELLOW  = "#ffd166"
RED     = "#ff6b6b"


# ── config ────────────────────────────────────────────────────────────────
CONFIG_PATH = Path.home() / ".setup_sheet_config.json"

# Built-in rules for setup sheet parsing (loaded into AI on first run, saved to config)
BUILTIN_RULES = [
    "Template selection: Excel column LASER → Laser sheet (doors). PUNCH or COMBI → Combi sheet (frames).",
    "Part Number: full identifier WITHOUT dimensions (e.g. ZZ9971904PXX_T1_EU_LH_F). Dimensions go in Description only.",
    "Doors (laser): T1 and T2 both mean SINGLE. Write 'SINGLE EURO DOOR LH', never 'TYPE 1' or 'TYPE 2'.",
    "Doors: EU=EURO (no 'LOCK'). STD=STANDARD LOCK. CL=CONLOCK. Only STD and CL get the word LOCK.",
    "Doors: F=FRONT, B=BACK, SET=both. RH=RIGHT HAND, LH=LEFT HAND.",
    "Frames (combi/punch): S=SINGLE, D=DOUBLE, F=FRAME. P=PICTURE (S_P_F/D_P_F = PICTURE FRAME). 3_SIDED/3SIDED/3-SIDED=3 SIDED. INTU/-INTU/INTUFRAME=INTUFRAME.",
    "Frames: strip _1OFF/_2OFF/_3OFF (and typos _10FF/_20FF) from Part Number. 1OFF/10FF → Pts/Prog=1, Pts/Sht=2; 2OFF/20FF → Pts/Prog=2, Pts/Sht=2; 3OFF → Pts/Prog=3, Pts/Sht=3.",
    "Description for doors: '1084X1774 SINGLE EURO DOOR LH' (dims + SINGLE + EURO + DOOR + hand).",
    "Description for frames: '850X2190 SINGLE FRAME RH' or '810X2000 SINGLE PICTURE FRAME LH' (S_P_F) or '790X1760 SINGLE FRAME LH INTUFRAME' or '1759X1960 DOUBLE DOOR EURO INTUFRAME' (D_F_INTU).",
    "Laser template uses machines L1030, L95, L49, L3030S, L20. Combi template uses TC1000, TC3000.",
    "NOTES field: when part code has _1OFF or _10FF, write '1OFF PROGRAM' (bold, 20pt). When _2OFF or _20FF, write '2OFF PROGRAM' (bold, 20pt).",
]

def load_config():
    cfg = {}
    if CONFIG_PATH.exists():
        try:
            cfg = json.loads(CONFIG_PATH.read_text())
        except Exception:
            pass
    # On first run or empty rules: load built-in rules so AI knows setup sheet conventions
    rules = cfg.get("learned_rules", [])
    if not rules:
        cfg["learned_rules"] = list(BUILTIN_RULES)
    # Ensure profiles exist; migrate flat config to Default if needed
    had_profiles = bool(cfg.get("profiles"))
    if "profiles" not in cfg or not cfg["profiles"]:
        d = _default_profile_data()
        for k in ["programmer", "sheet_size", "pts_prog", "pts_sht", "thickness", "material_type",
                  "tape", "inspect", "su_sheet", "last_excel", "last_laser_template",
                  "last_combi_template", "last_output", "start_row", "max_rows"]:
            if k in cfg:
                d[k] = cfg[k]
        cfg["profiles"] = {"Default": d}
    if "active_profile" not in cfg or cfg["active_profile"] not in cfg["profiles"]:
        cfg["active_profile"] = list(cfg["profiles"].keys())[0]
    if "verify_order_rules" not in cfg:
        cfg["verify_order_rules"] = []
    needs_save = not CONFIG_PATH.exists() or not rules or not had_profiles
    if needs_save:
        try:
            CONFIG_PATH.write_text(json.dumps(cfg, indent=2))
        except Exception:
            pass
    return cfg

def _default_profile_data():
    return {
        "programmer": "", "sheet_size": "2500x1250", "pts_prog": "1", "pts_sht": "1",
        "thickness": "1.2", "material_type": "ZINTEC",
        "tape": "NEW", "inspect": "FULLY", "su_sheet": "NEW",
        "last_excel": "", "last_laser_template": "", "last_combi_template": "",
        "last_output": "", "start_row": 1, "max_rows": 15,
        "excel_filter": "All", "max_scan": 500,
        "laser_ticks": [False] * 5, "combi_ticks": [False] * 2,
    }

def save_config(cfg: dict):
    CONFIG_PATH.write_text(json.dumps(cfg, indent=2))


# ── AI core — unified Anthropic + OpenAI wrapper ──────────────────────────

ANTHROPIC_MODELS = ["claude-sonnet-4-20250514", "claude-haiku-4-5-20251001"]
OPENAI_MODELS    = ["gpt-4o", "gpt-4o-mini", "gpt-4-turbo"]

def get_provider(cfg: dict) -> str:
    """Returns 'anthropic' or 'openai'."""
    return cfg.get("ai_provider", "anthropic")

def get_model(cfg: dict) -> str:
    provider = get_provider(cfg)
    if provider == "openai":
        return cfg.get("openai_model", "gpt-4o")
    return cfg.get("anthropic_model", "claude-sonnet-4-20250514")

def ai_call(system: str, user: str, cfg: dict,
            messages: list = None, max_tokens: int = 1000) -> str:
    """
    Unified AI call. Works with both Anthropic and OpenAI.
    If `messages` is provided (for chat), uses it directly.
    Otherwise wraps `user` as a single user message.
    """
    provider = get_provider(cfg)
    model    = get_model(cfg)

    if messages is None:
        messages = [{"role": "user", "content": user}]

    if provider == "openai":
        if not HAS_OPENAI:
            raise ImportError("openai package not installed. Run: pip install openai")
        key = cfg.get("openai_api_key") or os.environ.get("OPENAI_API_KEY", "")
        if not key:
            raise ValueError("No OpenAI API key found. Add it in Settings.")
        client   = _openai_lib.OpenAI(api_key=key)
        full_msgs = [{"role": "system", "content": system}] + messages
        resp = client.chat.completions.create(
            model=model, max_tokens=max_tokens, messages=full_msgs)
        return resp.choices[0].message.content.strip()

    else:  # anthropic
        if not HAS_ANTHROPIC:
            raise ImportError("anthropic package not installed. Run: pip install anthropic")
        key = cfg.get("api_key") or os.environ.get("ANTHROPIC_API_KEY", "")
        if not key:
            raise ValueError("No Anthropic API key found. Add it in Settings.")
        client = _anthropic_lib.Anthropic(api_key=key)
        resp = client.messages.create(
            model=model, max_tokens=max_tokens,
            system=system, messages=messages)
        return resp.content[0].text.strip()


def build_system_prompt(cfg: dict) -> str:
    """
    Build a system prompt that includes ALL setup sheet rules (hardcoded + user additions).
    This is injected into every AI call so the AI knows how setup sheets work and follows user rules.
    """
    # All hardcoded rules — the AI must follow these
    builtin = "\n".join(f"- {r}" for r in BUILTIN_RULES)
    part_rules = PART_CODE_RULES.strip()

    # User-added rules (from chat feedback or "Add rule") — these supplement or override
    learned = cfg.get("learned_rules", [])
    learned_text = ""
    if learned:
        learned_text = "\n\nUSER ADDITIONS (rules the user added — always follow these, they override if conflicting):\n"
        for i, r in enumerate(learned, 1):
            learned_text += f"  {i}. {r}\n"

    return f"""You are an expert manufacturing automation assistant for a CAD/CAM engineer.
You help automate laser cutting and punch press setup sheets by:
- Parsing part codes (like ZZ9978318931PXX_100X2000_T1_EU_LH_F)
- Analysing Excel job lists
- Analysing Word setup sheet templates
- Filling templates with the correct data

═══════════════════════════════════════════════════════════════
SETUP SHEET RULES (how the program works — follow these exactly)
═══════════════════════════════════════════════════════════════

{builtin}

PART CODE RULES (compact reference):
{part_rules}

TEMPLATE SELECTION (from Excel Laser/Punch column):
- LASER → Laser template (doors, machines L1030, L95, L49, L3030S, L20)
- PUNCH or COMBI → Combi template (frames, machines TC1000, TC3000)

TEMPLATE FIELDS (what goes where):
- customer, part_number, description, revision, dnc, programmer
- sheet_size, pts_prog, pts_sht, thickness, type (material_type)
- tape, inspect, su_sheet (dropdowns)

FIX REQUESTS: When the user asks to fix setup sheet(s), respond with:
- Single: FIX_REQUEST:{{"doc":"P44302","updates":{{"thickness":"2"}}}}
- Multiple: FIX_REQUEST:{{"docs":["P44300","P44301"],"updates":{{"pts_prog":"2"}}}}
- All: FIX_REQUEST:{{"docs":"ALL","updates":{{"thickness":"2"}}}}
- Field keys: pts_prog, pts_sht, thickness, type, sheet_size, customer, part_number, description, revision, dnc, programmer, tape, inspect, su_sheet
- Response MUST start with FIX_REQUEST:... then newline and brief confirmation.
{learned_text}
═══════════════════════════════════════════════════════════════
BEHAVIOUR:
- When the user asks you to ADD a new rule, REMEMBER something, or says "from now on..." — acknowledge it clearly. The system will save their rule automatically for future use.
- If the user reports a mistake or correction, acknowledge it and confirm what rule you will follow.
- Always be concise and practical."""


def ai_chat(messages: list, cfg: dict, log=None) -> str:
    """General chat with full conversation history."""
    return ai_call(build_system_prompt(cfg), "", cfg,
                   messages=messages, max_tokens=1500)


def ai_extract_rule(feedback: str, cfg: dict) -> str | None:
    """Distil feedback or 'add rule' request into a short persistent rule."""
    prompt = f"""The user gave feedback or asked to add a rule about setup sheet parsing or template filling:

\"{feedback}\"

Extract a SHORT, clear rule (1-2 sentences max) to remember going forward.
Examples: "When code has X, description = Y" or "INTU means INTUFRAME" or "Strip _1OFF from Part Number".
Return ONLY the rule text, nothing else. If no clear rule can be extracted, return empty string."""
    rule = ai_call("You extract clear, concise rules from user feedback or add-rule requests.", prompt,
                   cfg, max_tokens=150)
    return rule if rule else None


def ai_parse_code(part_code: str, cfg: dict, log) -> dict:
    log(f"  🤖 AI parsing: {part_code}")
    prompt = f"""Parse this part code and extract ALL meaningful fields.
Part code: {part_code}

Return ONLY valid JSON (no markdown):
{{
  "order_number": "...", "part_number": "...", "width": "...", "length": "...",
  "thickness": "...", "region": "...", "hand": "...", "finish": "...",
  "door_type": "...", "material": "...", "raw_code": "{part_code}", "notes": ""
}}"""
    raw = ai_call(build_system_prompt(cfg), prompt, cfg, max_tokens=500)
    raw = re.sub(r"^```json|^```|```$", "", raw, flags=re.MULTILINE).strip()
    parsed = json.loads(raw)
    log("  ✅ Parsed OK")
    return parsed


def ai_analyse_template(docx_path: str, cfg: dict, log) -> dict:
    log("  🤖 AI analysing Word template…")
    doc = Document(docx_path)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip(): lines.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip(): lines.append(cell.text.strip())

    prompt = f"""Analyse this manufacturing setup sheet Word template.
Text in document:
---
{chr(10).join(lines[:150])}
---
Identify EVERY field/label to fill in. Classify each as:
- "dynamic" = changes per job (order no, part no, dimensions)
- "fixed"   = same every sheet (sheet size, machine name, company)
- "date"    = today's date

Return ONLY valid JSON (no markdown):
{{
  "detected_placeholders": ["Order No:", "Part No:", "Sheet Size:", ...],
  "suggested_mapping": {{"Order No:": "order_number", "Sheet Size:": "sheet_size", ...}},
  "field_types": {{"Order No:": "dynamic", "Sheet Size:": "fixed", "Date:": "date"}},
  "notes": "brief description"
}}"""
    raw = ai_call(build_system_prompt(cfg), prompt, cfg, max_tokens=1000)
    raw = re.sub(r"^```json|^```|```$", "", raw, flags=re.MULTILINE).strip()
    result = json.loads(raw)
    log(f"  ✅ {result.get('notes','')}")
    log(f"  📋 {len(result.get('detected_placeholders',[]))} fields found")
    return result


def ai_analyse_excel(xlsx_path: str, cfg: dict, log) -> dict:
    log("  🤖 AI analysing Excel…")
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb.active
    rows_preview = []
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        if i > 30: break
        if any(v is not None for v in row):
            rows_preview.append([str(v) if v is not None else "" for v in row])

    prompt = f"""Analyse this manufacturing Excel (tab-separated, first 30 rows):
---
{chr(10).join([chr(9).join(r) for r in rows_preview])}
---
Return ONLY valid JSON:
{{"header_row":0,"part_code_column":"A","dnc_column":"B","other_columns":{{"C":"quantity"}},"notes":"..."}}"""
    raw = ai_call(build_system_prompt(cfg), prompt, cfg, max_tokens=500)
    raw = re.sub(r"^```json|^```|```$", "", raw, flags=re.MULTILINE).strip()
    result = json.loads(raw)
    log(f"  ✅ {result.get('notes','')}")
    return result


# ── ZZ order part code parser (regex-based) ─────────────────────────────────
def _is_frame_format(part_code: str) -> bool:
    """Frame format: _S_F, _D_F, _S_P_F, _D_P_F, _D_INTUFRAME, _3SIDED, etc. Door format: T1, T2, EU, STD, CL."""
    return bool(re.search(r"_([SD])(?:_P)?_F(?:_|-|$|\d)", part_code, re.I)
        or re.search(r"_([SD])_(?:INTU|INTUFRAME)(?:_|$)", part_code, re.I)
        or re.search(r"_3[-_ ]?SIDED", part_code, re.I))


def build_frame_description(part_code: str) -> str:
    """
    Build description for frame (combi/punch) codes.
    ZZ9971892PXX_850X2190_S_F_RH_2OFF → 850X2190 SINGLE FRAME RH
    ZZ9951681PEX_1739X2140_D_F → 1739X2140 DOUBLE FRAME
    ZZ9950967PEX_1514X2522_D_F_3_SIDED → 1514X2522 DOUBLE 3 SIDED FRAME
    ZZ9971758PXX_790X1760_S_F_LH_INTU → 790X1760 SINGLE FRAME LH INTUFRAME
    ZZ9971771PXX_1759X1960_D_F_INTU → 1759X1960 DOUBLE DOOR EURO INTUFRAME
    ZZ9971401PXX_810X2000_S_P_F_LH → 810X2000 SINGLE PICTURE FRAME LH
    """
    parts = re.split(r"\s*_AND_\s*", part_code, flags=re.I)
    descs = []
    count_map = {"S": "SINGLE", "D": "DOUBLE"}

    for p in parts:
        m = re.search(r"(\d+)[xX](\d+)", p)
        dims = f"{m.group(1)}X{m.group(2)}" if m else ""
        # S_F, D_F or S_P_F, D_P_F (P=PICTURE FRAME) or D_INTUFRAME (hyphen in D_F-INTU)
        m = re.search(r"_([SD])_P_F(?:_|$|\d)", p, re.I)
        is_picture = bool(m)
        if m:
            count = count_map.get((m.group(1) or "").upper(), "")
        else:
            m = re.search(r"_([SD])_F(?:_|-|$|\d)", p, re.I)
            if m:
                count = count_map.get((m.group(1) or "").upper(), "")
            else:
                m = re.search(r"_([SD])_(?:INTU|INTUFRAME)(?:_|$)", p, re.I)
                count = count_map.get((m.group(1) or "").upper(), "") if m else ""
        m = re.search(r"(?:^|_)(LH|RH)(?:_|$)", p, re.I)
        hand = (m.group(1) or "").upper() if m else ""
        # 3_SIDED, 3SIDED, 3 SIDED, 3-SIDED
        has_3sided = bool(re.search(r"_3[-_ ]?SIDED(?:_|$)", p, re.I))
        # INTU or -INTU or INTUFRAME (all = INTUFRAME)
        has_intu = bool(re.search(r"(?:[-_])?INTU(?:FRAME)?(?:_|$|[(\s])", p, re.I))

        # D_F_INTU (no hand) → DOUBLE DOOR EURO INTUFRAME
        if has_intu and count == "DOUBLE" and not hand:
            bits = [dims, "DOUBLE", "DOOR", "EURO", "INTUFRAME"]
        else:
            bits = [dims, count]
            if has_3sided:
                bits.append("3 SIDED")
            bits.append("PICTURE FRAME" if is_picture else "FRAME")
            if hand:
                bits.append(hand)
            if has_intu:
                bits.append("INTUFRAME")
        desc = " ".join(b for b in bits if b).strip()
        if desc:
            descs.append(desc)

    return "\n".join(descs) if descs else part_code


def build_door_description(part_code: str) -> str:
    """
    Build description for door (laser) codes.
    Format: "1084X1774 SINGLE EURO DOOR LH" (EU=EURO; only STD/CL get "LOCK")
    """
    parts = re.split(r"\s*_AND_\s*", part_code, flags=re.I)
    descs = []
    door_types = []
    region_map = {"EU": "EURO", "STD": "STANDARD LOCK", "CL": "CONLOCK", "US": "US", "UK": "UK"}

    for p in parts:
        m = re.search(r"(\d+)[xX](\d+)", p)
        dims = f"{m.group(1)}X{m.group(2)}" if m else ""
        m = re.search(r"T(\d+)", p, re.I)
        door_type = "SINGLE" if m else ""
        m = re.search(r"(?:^|_)(EU|STD|CL|US|UK)(?:_|$)", p, re.I)
        region = region_map.get((m.group(1) or "").upper(), (m.group(1) or "").upper()) if m else ""
        m = re.search(r"(?:^|_)(LH|RH)(?:_|$)", p, re.I)
        hand = (m.group(1) or "").upper() if m else ""
        m = re.search(r"_SET(?:_|$|\s|_AND)", p, re.I)
        fb = "SET" if m else ""
        if not fb:
            m = re.search(r"_([FB])(?:_|$|\s|_AND)", p, re.I)
            fb = (m.group(1) or "").upper() if m else ""
        door_types.append(fb)

        bits = [dims, door_type, region, "DOOR", hand]
        desc = " ".join(b for b in bits if b).strip()
        if desc:
            descs.append(desc)

    result = "\n".join(descs) if descs else part_code
    if len(parts) >= 2 and len(door_types) >= 2:
        if all(d == "B" for d in door_types):
            result = result + "\nCOMMON BACK"
        elif all(d == "F" for d in door_types):
            result = result + "\nCOMMON FRONT"
    return result


def build_description(part_code: str, is_frame: bool | None = None) -> str:
    """
    Build human-readable description. Auto-detects frame vs door format.
    Frame: _S_F or _D_F → build_frame_description
    Door: T1, EU, etc → build_door_description
    """
    if is_frame is None:
        is_frame = _is_frame_format(part_code)
    if is_frame:
        return build_frame_description(part_code)
    return build_door_description(part_code)


# ── Template selection (hardcoded, no AI) ────────────────────────────────────
def pick_template(laser_punch_col: str, laser_path: str, combi_path: str):
    """LASER → laser sheet; PUNCH/COMBI/else → combi sheet. Returns (is_laser, template_path)."""
    col = (laser_punch_col or "").upper()
    if "LASER" in col:
        return True, laser_path
    return False, combi_path


# ── Part code rules (for parsing and AI) ─────────────────────────────────────
PART_CODE_RULES = """
DOORS (LASER): T1/T2=SINGLE, EU=EURO, STD=STANDARD LOCK, CL=CONLOCK, F=FRONT, B=BACK, SET=both
FRAMES (PUNCH/COMBI): S=SINGLE, D=DOUBLE, F=FRAME, P=PICTURE (S_P_F/D_P_F=PICTURE FRAME)
3_SIDED/3SIDED/3-SIDED=3 SIDED. INTU/-INTU/INTUFRAME=INTUFRAME (all same)
RH=RIGHT HAND, LH=LEFT HAND
Strip _1OFF/_2OFF/_3OFF, _10FF/_20FF from Part Number. 1OFF/10FF→Pts/Prog=1,Pts/Sht=2; 2OFF/20FF→2,2; 3OFF→3,3
Part Number = identifier WITHOUT dimensions
Dimensions (NNNxNNNN) go in Description only
"""


def parse_zz_part_code(part_code: str) -> dict:
    """
    Parse ZZ order part codes. Part Number = full identifier e.g. ZZ9971905PXX_T1_EU_RH_F
    - ZZ9971906PXX_884X2224_T1_EU_RH_F
    - ZZ9971905PXX_T1_EU_RH_F_2OFF → part_number ZZ9971905PXX_T1_EU_RH_F (2OFF stripped, still sets Pts/Prog)
    - Multi: ZZ9971905PXX_884X2224_T1_EU_LH_B_AND_ZZ9971906PXX_884X2224_T1_EU_RH_B
    """
    def _str(v): return str(v) if v else ""
    result = {
        "order_number": "", "part_number": "", "width": "", "length": "",
        "dimensions": "", "thickness": "", "region": "", "hand": "", "door_type": "",
        "raw_code": part_code, "description": build_description(part_code)
    }
    # Split by _AND_ for multi-part
    parts = re.split(r"\s*_AND_\s*", part_code, flags=re.I)
    first = parts[0] if parts else ""
    # Part Number = full identifier WITHOUT dimensions (dims go in Description)
    # Also strip _1OFF/_2OFF for cleaner display (OFF suffix still drives Pts/Prog, Pts/Sht)
    # e.g. ZZ9971904PXX_1084X1774_T1_EU_LH_F → ZZ9971904PXX_T1_EU_LH_F
    # e.g. ZZ9971885PXX_1050X2190_S_F_LH_2OFF → ZZ9971885PXX_S_F_LH
    def _strip_dims(s):
        return re.sub(r"_\d+[xX]\d+_?", "_", s.strip()).rstrip("_")
    def _strip_off(s):
        # Strip _1OFF, _2OFF, _3OFF and typos _10FF, _20FF from Part Number
        s = re.sub(r"_[123]OFF(?:_|$)", "", s, flags=re.I).rstrip("_")
        s = re.sub(r"_10FF(?:_|$)", "", s, flags=re.I).rstrip("_")
        s = re.sub(r"_20FF(?:_|$)", "", s, flags=re.I).rstrip("_")
        return s
    if len(parts) == 1:
        result["part_number"] = _strip_off(_strip_dims(part_code))
    else:
        result["part_number"] = " / ".join(_strip_off(_strip_dims(p)) for p in parts)
    # ZZ + digits = order number
    m = re.search(r"ZZ(\d+)", first, re.I)
    if m:
        result["order_number"] = "ZZ" + m.group(1)
    # Dimensions: NNNxNNNN or NNNXNNNN
    m = re.search(r"(\d+)[xX](\d+)", first)
    if m:
        result["width"] = m.group(1)
        result["length"] = m.group(2)
        result["dimensions"] = f"{m.group(1)}x{m.group(2)}"
    # T1, T2 = thickness/type
    m = re.search(r"T(\d+)", first, re.I)
    if m:
        result["thickness"] = "T" + m.group(1)
    # Region: EU, STD, CL, US, UK
    m = re.search(r"(?:^|_)(EU|STD|CL|US|UK)(?:_|$)", first, re.I)
    if m:
        result["region"] = m.group(1).upper()
    # Hand: LH, RH
    m = re.search(r"(?:^|_)(LH|RH)(?:_|$)", first, re.I)
    if m:
        result["hand"] = m.group(1).upper()
    # Door/finish: F, B, SET (SET = both F&B)
    m = re.search(r"_SET(?:_|$|\s|_AND)", first, re.I)
    if m:
        result["door_type"] = "SET"
    else:
        m = re.search(r"_([FB])(?:_|$|\s|_AND)", first, re.I)
        if m:
            result["door_type"] = m.group(1).upper()
    return result


def read_register_excel(xlsx_path: str, start_row: int, max_rows: int, log,
                        filter_type: str = "", max_scan: int = 500) -> list:
    """
    Read Register.xlsx with fixed columns:
    A=Program Number (DNC), B=Link, C=Laser/Punch, D=Customer, E=Javelin, F=Part No, G=Rev
    - start_row: 0-based row to start (0=first row, 1=skip header)
    - max_rows: max data rows to return
    - filter_type: "" (all), "LASER", or "PUNCH" — only include matching rows
    - max_scan: stop after scanning this many rows (avoids reading thousands)
    """
    wb = openpyxl.load_workbook(xlsx_path, data_only=True, read_only=True)
    try:
        ws = wb["register"] if "register" in wb.sheetnames else wb.active
        # min_row/max_row are 1-based
        # When filtering (LASER/PUNCH), scan entire sheet—matching rows can be anywhere (e.g. row 22k+)
        min_row = start_row + 1
        if filter_type:
            max_row = None  # no limit: scan whole sheet to find filtered rows
        else:
            max_row = min_row + max_scan - 1 if max_scan > 0 else None
        rows = []
        for i, row in enumerate(ws.iter_rows(min_row=min_row, max_row=max_row, values_only=True)):
            if max_rows > 0 and len(rows) >= max_rows:
                break
            try:
                dnc = row[0]  # A
                link = row[1] if len(row) > 1 else dnc  # B
                laser_punch = (row[2] if len(row) > 2 else "")  # C
                customer = row[3] if len(row) > 3 else ""  # D
                part_code = row[5] if len(row) > 5 else ""  # F
                rev = row[6] if len(row) > 6 else "E"  # G, default E for ZZ
            except (IndexError, TypeError):
                continue
            if not part_code:
                continue
            lp_upper = _str(laser_punch).upper()
            # Apply filter: LASER only / PUNCH only
            if filter_type:
                if filter_type.upper() == "LASER" and "LASER" not in lp_upper:
                    continue
                if filter_type.upper() == "PUNCH" and "LASER" in lp_upper:
                    continue
            entry = {
                "dnc": _str(dnc) or _str(link),
                "laser_punch": lp_upper,
                "customer": _str(customer),
                "part_code": _str(part_code),
                "revision": _str(rev) or "E",
            }
            rows.append(entry)
        log(f"  📊 {len(rows)} rows read" + (f" (filter: {filter_type})" if filter_type else ""))
        return rows
    finally:
        wb.close()


def search_register_by_dnc(xlsx_path: str, dnc_query: str, log, max_scan: int = 100000) -> list:
    """
    Search the selected Excel (Register) file for DNC/Program Number across BOTH Laser and Punch.
    Searches ALL sheets. Matches ONLY column A (Program Number) - case-insensitive, EXACT match.
    Column B (Link) is ignored for search to avoid false matches.
    Returns list of matching rows (same format as read_register_excel).
    """
    if not dnc_query or not dnc_query.strip():
        return []
    q = _str(dnc_query).strip().upper()
    wb = openpyxl.load_workbook(xlsx_path, data_only=True, read_only=True)
    try:
        rows = []
        seen_dnc = set()
        total_scanned = 0
        # Search ALL sheets (Laser and Punch may be in different sheets)
        sheet_names = wb.sheetnames
        for sheet_name in sheet_names:
            if total_scanned >= max_scan:
                break
            ws = wb[sheet_name]
            count = 0
            for row in ws.iter_rows(min_row=2, values_only=True):
                if total_scanned >= max_scan:
                    break
                total_scanned += 1
                count += 1
                try:
                    dnc = _str(row[0]) if len(row) > 0 else ""
                    link = _str(row[1]) if len(row) > 1 else dnc
                    laser_punch = (row[2] if len(row) > 2 else "")
                    customer = row[3] if len(row) > 3 else ""
                    part_code = row[5] if len(row) > 5 else ""
                    rev = row[6] if len(row) > 6 else "E"
                except (IndexError, TypeError):
                    continue
                if not part_code:
                    continue
                dnc_upper = _str(dnc).strip().upper()
                if dnc_upper != q:
                    continue
                if dnc_upper in seen_dnc:
                    continue
                seen_dnc.add(dnc_upper)
                lp_upper = _str(laser_punch).upper()
                entry = {
                    "dnc": dnc or link,
                    "laser_punch": lp_upper,
                    "customer": _str(customer),
                    "part_code": _str(part_code),
                    "revision": _str(rev) or "E",
                }
                rows.append(entry)
        log(f"  🔍 Found {len(rows)} match(es) for '{dnc_query}' (scanned {total_scanned} rows across {len(sheet_names)} sheet(s))")
        return rows
    finally:
        wb.close()


def _str(v):
    return str(v).strip() if v is not None else ""


# DNC pattern: P followed by digits (e.g. P15783). Primary: 4-6 digits; fallback: 3-8 for OCR
_DNC_RE = re.compile(r"\bP\d{4,6}\b", re.IGNORECASE)
_DNC_RE_LOOSE = re.compile(r"\b[Pp]\s*\d{3,8}\b", re.IGNORECASE)


def _extract_dnc_from_text(text: str, loose: bool = False) -> str | None:
    """Search for DNC pattern in text. Returns first match or None. loose=True for OCR fallback."""
    if not text:
        return None
    m = (_DNC_RE_LOOSE if loose else _DNC_RE).search(text)
    if m:
        # Normalize: remove spaces, uppercase
        s = re.sub(r"\s", "", m.group(0)).upper()
        return s if re.match(r"P\d{3,8}$", s) else None
    return None


def _try_system_tesseract() -> bool:
    """If bundled tesseract failed, try system install. Returns True if configured."""
    if not HAS_PYTESSERACT:
        return False
    try:
        for pf in (os.environ.get("ProgramFiles", "C:\\Program Files"),
                   os.environ.get("ProgramFiles(x86)", "C:\\Program Files (x86)")):
            sys_path = os.path.join(pf, "Tesseract-OCR", "tesseract.exe")
            tessdata = os.path.join(pf, "Tesseract-OCR", "tessdata")
            if os.path.isfile(sys_path) and os.path.isdir(tessdata):
                pytesseract.pytesseract.tesseract_cmd = sys_path
                os.environ["TESSDATA_PREFIX"] = tessdata + os.sep
                return True
    except Exception:
        pass
    return False


def _extract_dnc_from_pdf_ocr(pdf_path: str) -> str | None:
    """
    Extract DNC from scanned PDF using OCR (pymupdf + pytesseract).
    Renders each page to image, runs OCR, searches for DNC pattern.
    Tries multiple PSM modes for better accuracy on forms/setup sheets.
    If bundled tesseract fails (e.g. missing DLLs), retries with system install.
    Returns None if OCR unavailable, fails, or no match.
    """
    if not HAS_PYMUPDF or not HAS_PYTESSERACT:
        return None

    def _run_ocr() -> str | None:
        doc = fitz.open(pdf_path)
        text_parts = []
        try:
            for page in doc:
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                img_bytes = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_bytes))
                for psm in ("6", "3"):
                    ocr_text = pytesseract.image_to_string(img, config=f"--psm {psm}")
                    if ocr_text:
                        text_parts.append(ocr_text)
        finally:
            doc.close()
        text = " ".join(text_parts)
        dnc = _extract_dnc_from_text(text, loose=False)
        if not dnc:
            dnc = _extract_dnc_from_text(text, loose=True)
        return dnc

    try:
        return _run_ocr()
    except Exception:
        if _try_system_tesseract():
            try:
                return _run_ocr()
            except Exception:
                pass
        return None


def extract_dnc_from_pdf(pdf_path: str) -> str | None:
    """
    Extract DNC/Program number from PDF. Returns first match of P followed by 4–6 digits.
    Tries text extraction first (digital PDFs), then OCR fallback (scanned PDFs).
    Returns None if no text or no match. Requires pypdf; OCR requires pymupdf + pytesseract + Tesseract.
    """
    if not HAS_PYPDF:
        return None
    # 1. Try text extraction (works for digital PDFs)
    try:
        reader = PdfReader(pdf_path)
        text_parts = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
        text = " ".join(text_parts)
        dnc = _extract_dnc_from_text(text)
        if dnc:
            return dnc
    except Exception:
        pass
    # 2. Fallback: OCR for scanned/image-based PDFs
    return _extract_dnc_from_pdf_ocr(pdf_path)


# ZZ part number pattern (e.g. ZZ9972060PXX from order or ZZ9972060PXX_274X1187_T1_EU_RH_F in register)
_ZZ_RE = re.compile(r"\bZZ\d{6,}P[A-Z0-9]+", re.IGNORECASE)


def _extract_zz_base(part_code: str) -> str | None:
    """
    Extract ZZ base (e.g. ZZ9972060PXX) from full part code.
    Handles: ZZ9972060PXX_274X1187_T1_EU_RH_F → returns ZZ9972060PXX (ignores rest after _).
    """
    if not part_code:
        return None
    s = str(part_code).strip()
    m = _ZZ_RE.search(s)
    if m:
        base = m.group(0).upper()
        if "_" in base:
            base = base.split("_")[0]
        return base
    return None


def extract_pdf_header_for_verify(pdf_path: str, full_text: str) -> str:
    """
    Extract order header / YourRef section from PDF text for use in part parsing.
    YourRef often contains TYPE N I LH etc. Returns first ~40 lines to capture header.
    """
    if not full_text:
        return ""
    lines = full_text.replace("\r", "\n").split("\n")
    header_lines = []
    for i, line in enumerate(lines):
        if i >= 45:
            break
        header_lines.append(line)
        if "YourRef" in line or "Your Ref" in line:
            for j in range(i + 1, min(i + 5, len(lines))):
                header_lines.append(lines[j])
            break
    return " ".join(header_lines)


def extract_zz_with_descriptions(pdf_path: str) -> list[tuple[str, str, str]]:
    """
    Extract ZZ numbers with their description line from order PDF.
    Returns list of (zz_base, description, chunk). Chunk includes PDF header (YourRef etc) + line context.
    """
    text = ""
    if HAS_PYPDF:
        try:
            reader = PdfReader(pdf_path)
            parts = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
            text = "\n".join(parts)
        except Exception:
            pass
    if not text or not re.search(_ZZ_RE, text):
        text = ""
        if HAS_PYMUPDF and HAS_PYTESSERACT:
            try:
                doc = fitz.open(pdf_path)
                for page in doc:
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                    img = Image.open(io.BytesIO(pix.tobytes("png")))
                    for psm in ("6", "3"):
                        ocr_text = pytesseract.image_to_string(img, config=f"--psm {psm}")
                        if ocr_text:
                            text += "\n" + ocr_text
                doc.close()
            except Exception:
                pass
    header = extract_pdf_header_for_verify(pdf_path, text)
    result = []
    lines = text.replace("\r", "\n").split("\n")
    for i, line in enumerate(lines):
        for m in _ZZ_RE.finditer(line):
            base = m.group(0).upper()
            if "_" in base:
                base = base.split("_")[0]
            desc = ""
            line_chunk = " ".join(lines[i : i + 20])
            chunk = (header + " " + line_chunk) if header else line_chunk
            if "door" in chunk.lower():
                desc = "door"
            elif "frame" in chunk.lower():
                desc = "frame"
            result.append((base, desc, chunk))
    return result


def _is_door_from_description(desc: str) -> bool:
    """Door → LASER, Frame → PUNCH. Returns True for door (LASER)."""
    d = (desc or "").lower()
    if "frame" in d and "door" not in d:
        return False  # PUNCH
    if "door" in d or "door leaf" in d:
        return True  # LASER
    if "frame" in d:
        return False  # PUNCH
    return True  # default to LASER if unclear


# Verify Order: build Part No from order text. Dimensions: NNNxNNNN. Doors: A1/A2→T1, RH/LH, Euro→EU, Std→STD, Concealed→CL.
# Frames: TYPE N = Single, TYPE N I = Single+Intuframe, I = Intuframe, LH/RH/DR = hand.
_DIM_RE = re.compile(r"(\d{3,4})\s*[xX]\s*(\d{3,4})")
_VERIFY_DOOR_RULES = [
    (r"\bA1\b|\bA2\b", "T1"),
    (r"\bEuro\s*lock\b|Euro\s*locks?|Eurolock", "EU"),
    (r"\bStandard\b|Std\b", "STD"),
    (r"\bConcealed\s*lock\b|Concealed\s*locks?|Conlock|CL\b", "CL"),
]
_VERIFY_FRAME_RULES = [
    (r"\bSingle\s*frame\b|TYPENI\s*[LR]H|TYPE\s*N\s*I\s*[LR]H|TYPE\s*N\s*I\s*DR", "S_F"),
    (r"\bDouble\s*frame\b|TYPE\s*N\s*I[-]?\s*DR\b|I-\s*DR|TYPENI\s*DR", "D_F"),
]
_VERIFY_FRAME_EXTRA = [
    (r"\bIntuframe\b|Intu\s*frame|TYPE\s*N\s*I\b|TYPENI", "INTU"),
    # Eurolock|EU removed — frames don't add EU by default. Add custom rule "Eurolock|EU" if needed.
]


def _parse_custom_verify_rules(rules: list[str]) -> list[tuple[str, str]]:
    """Parse 'pattern|replacement' strings into (pattern, replacement) pairs."""
    out = []
    for r in rules or []:
        r = (r or "").strip()
        if "|" in r:
            pat, repl = r.split("|", 1)
            pat, repl = pat.strip(), repl.strip()
            if pat and repl:
                out.append((pat, repl))
    return out


def build_part_description_from_order(zz: str, chunk: str, is_door: bool, custom_rules: list[str] | None = None) -> str:
    """
    Build Part No from order text. E.g. ZZ9971784PXX_774x1921_T1_EU_RH (user adds F/B/SET).
    Returns base description without F/B/SET suffix.
    custom_rules: list of "pattern|replacement" — when pattern matches, replacement is added as segment.
    """
    chunk_norm = " " + (chunk or "").replace("\n", " ") + " "
    dim_match = _DIM_RE.search(chunk_norm)
    dims = ""
    if dim_match:
        dims = f"{dim_match.group(1)}x{dim_match.group(2)}"

    parsed_custom = _parse_custom_verify_rules(custom_rules or [])

    if is_door:
        lock = "EU"
        if re.search(r"\bConcealed\s*lock\b|Concealed\s*locks?|Conlock|CL\b", chunk_norm, re.I):
            lock = "CL"
        elif re.search(r"\bStandard\b|Std\b", chunk_norm, re.I):
            lock = "STD"
        elif re.search(r"\bEuro\s*lock\b|Euro\s*locks?|Eurolock", chunk_norm, re.I):
            lock = "EU"
        hand = ""
        if re.search(r"\bRH\b", chunk_norm):
            hand = "RH"
        elif re.search(r"\bLH\b", chunk_norm):
            hand = "LH"
        if not hand and re.search(r"\bA2\b", chunk_norm):
            hand = "RH"
        elif not hand and re.search(r"\bA1\b", chunk_norm):
            hand = "LH"
        extra = []
        for pat, repl in parsed_custom:
            try:
                if re.search(pat, chunk_norm, re.I):
                    extra.append(repl)
            except re.error:
                pass
        if dims and hand:
            base = f"{zz}_{dims}_T1_{lock}_{hand}"
            return base + ("_" + "_".join(extra) if extra else "")
        if dims:
            return f"{zz}_{dims}"
    else:
        frame_type = ""
        for pat, val in _VERIFY_FRAME_RULES:
            try:
                if re.search(pat, chunk_norm, re.I):
                    frame_type = val
                    break
            except re.error:
                pass
        hand = ""
        if re.search(r"\bRH\b", chunk_norm):
            hand = "_RH"
        elif re.search(r"\bLH\b", chunk_norm):
            hand = "_LH"
        elif re.search(r"\bDR\b", chunk_norm):
            hand = "_DR"
        extra_segments = []
        for pat, val in _VERIFY_FRAME_EXTRA:
            try:
                if re.search(pat, chunk_norm, re.I):
                    extra_segments.append(val)
            except re.error:
                pass
        for pat, repl in parsed_custom:
            try:
                if re.search(pat, chunk_norm, re.I):
                    extra_segments.append(repl)
            except re.error:
                pass
        extra_str = "_" + "_".join(extra_segments) if extra_segments else ""
        if dims and frame_type:
            return f"{zz}_{dims}_{frame_type}{hand}{extra_str}"
        if dims:
            return f"{zz}_{dims}"
    return zz


def extract_customer_from_pdf(pdf_path: str) -> str:
    """Extract customer name from order PDF. Returns uppercase (e.g. ASPEX)."""
    text = ""
    if HAS_PYPDF:
        try:
            reader = PdfReader(pdf_path)
            text = "\n".join(p.extract_text() or "" for p in reader.pages)
        except Exception:
            pass
    lines = text.split("\n")
    for i, line in enumerate(lines):
        line_lower = line.lower()
        if "to deliver to" in line_lower and i + 1 < len(lines):
            return lines[i + 1].strip().upper()
        if "aspex" in line_lower:
            return "ASPEX"
        if "mason" in line_lower and "king" in line_lower:
            return "MASON & KING"
    return "ASPEX"


def extract_zz_numbers_from_pdf(pdf_path: str) -> list[str]:
    """
    Extract all ZZ part numbers from PDF (order sheet). Tries text extraction first, then OCR.
    Returns sorted unique list of ZZ bases (e.g. ZZ9971947PXX).
    """
    text = ""
    if HAS_PYPDF:
        try:
            reader = PdfReader(pdf_path)
            parts = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
            text = " ".join(parts)
        except Exception:
            pass
    if not text or not re.search(_ZZ_RE, text):
        text = ""
        if HAS_PYMUPDF and HAS_PYTESSERACT:
            try:
                doc = fitz.open(pdf_path)
                for page in doc:
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                    img = Image.open(io.BytesIO(pix.tobytes("png")))
                    for psm in ("6", "3"):
                        ocr_text = pytesseract.image_to_string(img, config=f"--psm {psm}")
                        if ocr_text:
                            text += " " + ocr_text
                doc.close()
            except Exception:
                pass
    found = set()
    for m in _ZZ_RE.finditer(text):
        s = m.group(0).upper()
        if "_" in s:
            s = s.split("_")[0]
        found.add(s)
    return sorted(found)


_DNC_NUM_RE = re.compile(r"P(\d+)", re.IGNORECASE)


def get_last_dnc_for_type(xlsx_path: str, lp_type: str, sheet_index: int = 0) -> str:
    """
    Get next DNC number for LASER or PUNCH. Filters by Laser/Punch column (C), finds max DNC, increments.
    Returns e.g. P44305 or P15786.
    """
    if not xlsx_path or not os.path.isfile(xlsx_path):
        return "P1"
    wb = openpyxl.load_workbook(xlsx_path, data_only=True, read_only=True)
    try:
        sheet_name = wb.sheetnames[sheet_index] if sheet_index < len(wb.sheetnames) else wb.sheetnames[0]
        ws = wb[sheet_name]
        max_num = 0
        for row in ws.iter_rows(min_row=2, values_only=True):
            try:
                dnc = _str(row[0]) if len(row) > 0 else ""
                lp = _str(row[2]) if len(row) > 2 else ""
            except (IndexError, TypeError):
                continue
            if lp.upper() != lp_type.upper():
                continue
            m = _DNC_NUM_RE.search(dnc)
            if m:
                n = int(m.group(1))
                if n > max_num:
                    max_num = n
        return f"P{max_num + 1}"
    finally:
        wb.close()


def add_row_to_register_with_dnc(xlsx_path: str, laser_punch: str, customer: str,
                        part_no: str, rev: str = "E", sheet_index: int = 0) -> tuple[bool, str]:
    """
    Load Excel once, compute next DNC, insert row, save. Returns (success, dnc_used).
    Avoids separate get_last_dnc + add_row which would load the file twice.
    """
    if not xlsx_path or not os.path.isfile(xlsx_path):
        return False, ""
    try:
        from openpyxl.styles import PatternFill, Font
        red_fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
        red_font = Font(color="FF0000")
        wb = openpyxl.load_workbook(xlsx_path)
        sheet_name = wb.sheetnames[sheet_index] if sheet_index < len(wb.sheetnames) else wb.sheetnames[0]
        ws = wb[sheet_name]
        max_num = 0
        for row in ws.iter_rows(min_row=2, values_only=True):
            try:
                dnc = _str(row[0]) if len(row) > 0 else ""
                lp = _str(row[2]) if len(row) > 2 else ""
            except (IndexError, TypeError):
                continue
            if lp.upper() != laser_punch.upper():
                continue
            m = _DNC_NUM_RE.search(dnc)
            if m:
                n = int(m.group(1))
                if n > max_num:
                    max_num = n
        dnc = f"P{max_num + 1}"
        ws.insert_rows(2, 1)
        dnc_cell = ws.cell(row=2, column=1, value=dnc)
        dnc_cell.fill = red_fill
        dnc_cell.font = red_font
        ws.cell(row=2, column=2, value=dnc)
        ws.cell(row=2, column=3, value=laser_punch)
        ws.cell(row=2, column=4, value=customer)
        ws.cell(row=2, column=5, value="")
        ws.cell(row=2, column=6, value=part_no)
        ws.cell(row=2, column=7, value=rev)
        if ws.auto_filter:
            ws.auto_filter = None
        wb.save(xlsx_path)
        wb.close()
        return True, dnc
    except Exception:
        return False, ""


def add_row_to_register(xlsx_path: str, dnc: str, laser_punch: str, customer: str,
                        part_no: str, rev: str = "E", sheet_index: int = 0) -> bool:
    """
    Insert new row at row 2 in register. Columns: A=DNC, B=Link (same as DNC), C=Laser/Punch,
    D=Customer, E=Javelin (empty), F=Part No, G=Rev. Applies red background to DNC cell (column A).
    """
    if not xlsx_path or not os.path.isfile(xlsx_path):
        return False
    try:
        from openpyxl.styles import PatternFill, Font
        red_fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
        red_font = Font(color="FF0000")
        wb = openpyxl.load_workbook(xlsx_path)
        sheet_name = wb.sheetnames[sheet_index] if sheet_index < len(wb.sheetnames) else wb.sheetnames[0]
        ws = wb[sheet_name]
        ws.insert_rows(2, 1)
        dnc_cell = ws.cell(row=2, column=1, value=dnc)
        dnc_cell.fill = red_fill
        dnc_cell.font = red_font
        ws.cell(row=2, column=2, value=dnc)
        ws.cell(row=2, column=3, value=laser_punch)
        ws.cell(row=2, column=4, value=customer)
        ws.cell(row=2, column=5, value="")
        ws.cell(row=2, column=6, value=part_no)
        ws.cell(row=2, column=7, value=rev)
        # Clear autofilter so Excel shows correct data when reopened (avoids stale filter state)
        if ws.auto_filter:
            ws.auto_filter = None
        wb.save(xlsx_path)
        wb.close()
        return True
    except Exception:
        return False


def get_register_zz_bases(xlsx_path: str, max_scan: int = 100000) -> set[str]:
    """
    Get all ZZ part number bases from Register Excel.
    Scans the ENTIRE file: ALL sheets (PUNCH, LASER, etc.), ALL columns, ALL rows.
    Full part codes like ZZ9972060PXX_274X1187_T1_EU_RH_F → extracts ZZ9972060PXX.
    Returns set of uppercase ZZ bases.
    """
    if not xlsx_path or not os.path.isfile(xlsx_path):
        return set()
    wb = openpyxl.load_workbook(xlsx_path, data_only=True, read_only=True)
    bases = set()
    try:
        total = 0
        for sheet_name in wb.sheetnames:
            if total >= max_scan:
                break
            ws = wb[sheet_name]
            for row in ws.iter_rows(min_row=1, values_only=True):
                if total >= max_scan:
                    break
                total += 1
                for cell in row:
                    try:
                        part_code = _str(cell) if cell is not None else ""
                    except (IndexError, TypeError):
                        continue
                    base = _extract_zz_base(part_code)
                    if base:
                        bases.add(base)
        return bases
    finally:
        wb.close()


def build_dnc_to_part_code(xlsx_path: str, max_scan: int = 100000) -> dict[str, str]:
    """
    Build a mapping DNC -> part_code from the Register Excel.
    Uses column A (DNC) and column F (Part No). First match wins per DNC.
    """
    if not xlsx_path or not os.path.isfile(xlsx_path):
        return {}
    wb = openpyxl.load_workbook(xlsx_path, data_only=True, read_only=True)
    mapping = {}
    try:
        total = 0
        for sheet_name in wb.sheetnames:
            if total >= max_scan:
                break
            ws = wb[sheet_name]
            for row in ws.iter_rows(min_row=2, values_only=True):
                if total >= max_scan:
                    break
                total += 1
                try:
                    dnc = _str(row[0]) if len(row) > 0 else ""
                    part_code = _str(row[5]) if len(row) > 5 else ""
                except (IndexError, TypeError):
                    continue
                if not dnc or not part_code:
                    continue
                dnc_upper = dnc.upper()
                if dnc_upper not in mapping:
                    mapping[dnc_upper] = part_code
        return mapping
    finally:
        wb.close()


def sanitize_filename(s: str) -> str:
    """Replace invalid filename chars with underscore."""
    return re.sub(r'[/\\:*?"<>|]', "_", str(s).strip())


# ── template filling: write to value cells + content controls ─────────────
from docx.oxml.ns import qn
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"


def _get_cell_at(table, row_idx: int, col_idx: int):
    """Get cell at (row, col). Uses table.cell() which handles grid. Returns (cell, None) or (None, None)."""
    try:
        cell = table.cell(row_idx, col_idx)
        return cell, None
    except Exception:
        return None, None


def _set_cell_text(cell, text: str, font_pt: int = 14):
    """Set cell text, replacing existing content. Handles \\n as line breaks. Black, readable size."""
    if cell is None or not HAS_DOCX:
        return
    text = str(text)
    # Clear existing
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    if cell.paragraphs:
        p = cell.paragraphs[0]
        p.clear()
        lines = text.split("\n")
        for i, line in enumerate(lines):
            if i > 0:
                p.add_run().add_break()
            r = p.add_run(line)
            r.font.size = Pt(font_pt)
            r.font.color.rgb = RGBColor(0, 0, 0)
    else:
        p = cell.add_paragraph()
        for i, line in enumerate(text.split("\n")):
            if i > 0:
                p.add_run().add_break()
            r = p.add_run(line)
            r.font.size = Pt(font_pt)
            r.font.color.rgb = RGBColor(0, 0, 0)


def _set_cell_text_bold(cell, text: str, font_pt: int = 20):
    """Set cell text with bold and font size. Black, bold, for NOTES field."""
    if cell is None or not HAS_DOCX:
        return
    text = str(text)
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    if cell.paragraphs:
        p = cell.paragraphs[0]
        p.clear()
        lines = text.split("\n")
        for i, line in enumerate(lines):
            if i > 0:
                p.add_run().add_break()
            r = p.add_run(line)
            r.font.size = Pt(font_pt)
            r.font.color.rgb = RGBColor(0, 0, 0)
            r.font.bold = True
    else:
        p = cell.add_paragraph()
        for i, line in enumerate(text.split("\n")):
            if i > 0:
                p.add_run().add_break()
            r = p.add_run(line)
            r.font.size = Pt(font_pt)
            r.font.color.rgb = RGBColor(0, 0, 0)
            r.font.bold = True


WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"


def _find_notes_content_target(doc):
    """Find where to write NOTES content. Returns (target_element, 'tc'|'run').
    - Table: tc (table cell) element
    - Text box: the w:r (run) element after 'NOTES:' in wps:txbx"""
    try:
        # 1. Table cells (original logic)
        for tbl in doc.element.body.iter(qn("w:tbl")):
            for tr in tbl.iter(qn("w:tr")):
                tcs = [el for el in tr if el.tag == qn("w:tc")]
                for i, tc in enumerate(tcs):
                    text = "".join(t.text or "" for t in tc.iter(qn("w:t"))).strip().upper()
                    if "NOTES:" in text or text == "NOTES":
                        if i + 1 < len(tcs):
                            return (tcs[i + 1], "tc")
        # 2. Text box (wps:txbx) - NOTES: in one run, content in next run
        for txbx in doc.element.body.iter():
            if txbx.tag != "{%s}txbx" % WPS_NS:
                continue
            for txbx_content in txbx.iter(qn("w:txbxContent")):
                for p in txbx_content.findall(qn("w:p")):
                    runs = list(p.findall(qn("w:r")))
                    for j, r in enumerate(runs):
                        for t in r.iter(qn("w:t")):
                            if t.text and ("NOTES:" in t.text.upper() or t.text.strip().upper() == "NOTES"):
                                if j + 1 < len(runs):
                                    return (runs[j + 1], "run")
                                break
    except Exception:
        pass
    return (None, None)


def _set_tc_text_bold(tc, text: str, font_pt: int = 20):
    """Set text in a tc (table cell) XML element with bold and font size. Avoids Cell API."""
    if tc is None or not HAS_DOCX:
        return
    from docx.oxml import OxmlElement
    # Clear existing paragraphs
    for p in tc.findall(qn("w:p")):
        tc.remove(p)
    # Add new paragraph with run (centered)
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)
    p.append(pPr)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    b = OxmlElement("w:b")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(font_pt * 2))  # half-points
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(font_pt * 2))
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rPr.append(b)
    rPr.append(sz)
    rPr.append(szCs)
    rPr.append(color)
    t = OxmlElement("w:t")
    t.text = str(text)
    r.append(rPr)
    r.append(t)
    p.append(r)
    tc.append(p)


def _set_run_text_bold(run_elem, text: str, font_pt: int = 20):
    """Set text in a w:r (run) element with bold and font size. NOTES: stays left, content centered."""
    if run_elem is None or not HAS_DOCX:
        return
    from docx.oxml import OxmlElement
    p = run_elem.getparent()
    if p is not None and p.tag == qn("w:p"):
        # Add center tab stop so "2OFF PROGRAM" centers; NOTES: stays left
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p.insert(0, pPr)
        # Remove any paragraph-level center (jc) so NOTES: stays left
        jc = pPr.find(qn("w:jc"))
        if jc is not None:
            pPr.remove(jc)
        # Add center tab at ~3.5" (5040 twips) so text after tab is centered
        tabs = pPr.find(qn("w:tabs"))
        if tabs is None:
            tabs = OxmlElement("w:tabs")
            pPr.insert(0, tabs)
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "center")
        tab.set(qn("w:pos"), "5040")
        tabs.append(tab)
        # Insert tab run before our content run (NOTES: | tab | 2OFF PROGRAM)
        tab_run = OxmlElement("w:r")
        tab_run.append(OxmlElement("w:tab"))
        run_elem.addprevious(tab_run)
    # Clear existing w:t and set new content
    for t in run_elem.findall(qn("w:t")):
        run_elem.remove(t)
    # Add or update rPr (run properties)
    rPr = run_elem.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run_elem.insert(0, rPr)
    # Clear and set bold, size, color
    for child in list(rPr):
        rPr.remove(child)
    b = OxmlElement("w:b")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(font_pt * 2))
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(font_pt * 2))
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rPr.append(b)
    rPr.append(sz)
    rPr.append(szCs)
    rPr.append(color)
    t = OxmlElement("w:t")
    t.text = str(text)
    run_elem.append(t)


def _set_sdt_dropdown(doc, value: str):
    """Set first matching dropdown content control to value. Walks all elements."""
    for body in doc.element.body.iter():
        if body.tag != qn("w:sdt"):
            continue
        sdtPr = body.find(qn("w:sdtPr"))
        if sdtPr is None:
            continue
        combo = sdtPr.find(qn("w:comboBox"))
        if combo is None:
            continue
        # Check if value is in list
        vals = [li.get(qn("w:value")) for li in combo.findall(qn("w:listItem")) if li.get(qn("w:value"))]
        if value not in vals and value:
            continue  # try next dropdown
        sdtContent = body.find(qn("w:sdtContent"))
        if sdtContent is None:
            continue
        for t in sdtContent.iter(qn("w:t")):
            if t.text is not None:
                t.text = value or "Choose an item."
                return True
        # No w:t yet, add one
        for r in sdtContent.iter(qn("w:r")):
            t = r.find(qn("w:t"))
            if t is not None:
                t.text = value or "Choose an item."
                return True
        # Add run with text
        r = type(body)("w:r", nsmap=body.nsmap)
        t = type(body)("w:t", nsmap=body.nsmap)
        t.text = value or "Choose an item."
        r.append(t)
        sdtContent.append(r)
        return True
    return False


def _set_run_black_bold(run_elem, font_size_half_pt: int = 28):
    """Set run to black color and readable font size. font_size_half_pt: 24=12pt, 28=14pt."""
    from docx.oxml import OxmlElement
    rPr = run_elem.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run_elem.insert(0, rPr)
    # Black color - set val, remove theme attrs so it's actual black
    color = rPr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        rPr.append(color)
    color.set(qn("w:val"), "000000")
    # Font size (half-points)
    for tag in ("w:sz", "w:szCs"):
        sz = rPr.find(qn(tag))
        if sz is None:
            sz = OxmlElement(tag)
            sz.set(qn("w:val"), str(font_size_half_pt))
            rPr.append(sz)
        else:
            sz.set(qn("w:val"), str(font_size_half_pt))


def _set_sdt_dropdown_by_order(doc, index: int, value: str):
    """Set Nth dropdown (0=Tape, 1=Inspect, 2=SU Sheet) to value. Black text, 12pt."""
    count = 0
    for body in doc.element.iter():
        if body.tag != qn("w:sdt"):
            continue
        sdtPr = body.find(qn("w:sdtPr"))
        if sdtPr is None:
            continue
        combo = sdtPr.find(qn("w:comboBox"))
        if combo is None:
            continue
        if count == index:
            sdtContent = body.find(qn("w:sdtContent"))
            if sdtContent is not None:
                for t in sdtContent.iter(qn("w:t")):
                    if t.text is not None:
                        t.text = value or "Choose an item."
                        run = t.getparent()
                        if run is not None and run.tag == qn("w:r"):
                            _set_run_black_bold(run, 28)  # 14pt
                        return True
                for r in sdtContent.iter(qn("w:r")):
                    t = r.find(qn("w:t"))
                    if t is not None:
                        t.text = value or "Choose an item."
                        _set_run_black_bold(r, 28)  # 14pt
                        return True
            return False
        count += 1
    return False


def _set_sdt_checkbox_by_order(doc, index: int, checked: bool):
    """Set Nth checkbox (0=L1030, 1=L95, 2=L49, 3=L3030S, 4=L20) to checked state."""
    count = 0
    for body in doc.element.iter():
        if body.tag != qn("w:sdt"):
            continue
        sdtPr = body.find(qn("w:sdtPr"))
        if sdtPr is None:
            continue
        cb = sdtPr.find(qn("w14:checkbox"))
        if cb is None:
            continue
        if count == index:
            chk = cb.find(qn("w14:checked"))
            if chk is not None:
                chk.set(qn("w14:val"), "1" if checked else "0")
            # Update w:sym display char: checked=00FE, unchecked=00A8 (Wingdings)
            sdtContent = body.find(qn("w:sdtContent"))
            if sdtContent is not None:
                for sym in sdtContent.iter(qn("w:sym")):
                    sym.set(qn("w:char"), "00FE" if checked else "00A8")
                    break
            return True
        count += 1
    return False


# Laser template: (row, col) for VALUE cells (from _inspect.py)
LASER_FIELDS = {
    "customer": (0, 1), "part_number": (0, 7), "description": (2, 1),
    "revision": (2, 7), "assy_number": (4, 7), "dnc": (6, 1), "programmer": (6, 7),
    "sheet_size": (8, 1), "pts_prog": (8, 7), "pts_sht": (8, 10),
    "thickness": (10, 7), "type": (10, 10),
}
# Combi template
COMBI_FIELDS = {
    "customer": (0, 2), "part_number": (0, 9), "description": (2, 2),
    "revision": (2, 9), "assy_number": (4, 9), "dnc": (6, 2), "programmer": (6, 9),
    "sheet_size": (8, 2), "thickness": (8, 8), "type": (8, 12),
    "pts_prog": (10, 8), "pts_sht": (10, 12),
}


def fill_word_template(template_path, output_path, data, opts, log):
    """
    Fill template with data. opts = dict with:
    - is_laser: bool (Laser vs Combi template)
    - sheet_size, pts_prog, pts_sht, thickness, material_type: from user
    - programmer: from settings
    - tape, inspect, su_sheet: dropdown values (NEW, FULLY, NEW etc)
    - laser_ticks: list of 5 bools [L1030, L95, L49, L3030S, L20] (Laser template)
    - combi_ticks: list of 2 bools [TC1000, TC3000] (Combi template)
    """
    doc = Document(template_path)
    fields = LASER_FIELDS if opts.get("is_laser", True) else COMBI_FIELDS
    table = doc.tables[0]

    # Build value dict from data + opts
    pts_prog = opts.get("pts_prog", "")
    pts_sht = opts.get("pts_sht", "")
    # Override from _1OFF/_2OFF/_3OFF (and typos _10FF/_20FF) in part code
    # 1OFF/10FF → Pts/Prog=1, Pts/Sht=2; 2OFF/20FF → Pts/Prog=2, Pts/Sht=2; 3OFF → Pts/Prog=3, Pts/Sht=3
    part_code = data.get("part_code", "") or data.get("part_number", "") or data.get("raw_code", "")
    off_match = re.search(r"_3OFF(?:_|$)", part_code, re.I)
    if off_match:
        pts_prog, pts_sht = "3", "3"
    else:
        off_match = re.search(r"_20FF(?:_|$)", part_code, re.I)  # typo for 2OFF
        if off_match:
            pts_prog, pts_sht = "2", "2"
        else:
            off_match = re.search(r"_10FF(?:_|$)", part_code, re.I)  # typo for 1OFF
            if off_match:
                pts_prog, pts_sht = "1", "2"
            else:
                off_match = re.search(r"_([12])OFF(?:_|$)", part_code, re.I)
                if off_match:
                    n = off_match.group(1)
                    pts_prog = "1" if n == "1" else "2"
                    pts_sht = "2"

    # NOTES field: 1OFF/10FF → "1OFF PROGRAM", 2OFF/20FF → "2OFF PROGRAM" (bold, 20pt)
    notes_text = ""
    if re.search(r"_20FF(?:_|$)", part_code, re.I) or re.search(r"_2OFF(?:_|$)", part_code, re.I):
        notes_text = "2OFF PROGRAM"
    elif re.search(r"_10FF(?:_|$)", part_code, re.I) or re.search(r"_1OFF(?:_|$)", part_code, re.I):
        notes_text = "1OFF PROGRAM"

    values = {
        "customer": data.get("customer", ""),
        "part_number": data.get("part_number", ""),
        "description": data.get("description", data.get("part_code", "")),
        "revision": data.get("revision", "E"),
        "assy_number": "",  # blank for ZZ orders
        "dnc": data.get("dnc", ""),
        "programmer": opts.get("programmer", ""),
        "sheet_size": opts.get("sheet_size", ""),
        "pts_prog": pts_prog,
        "pts_sht": pts_sht,
        "thickness": opts.get("thickness", data.get("thickness", "")),
        "type": opts.get("material_type", data.get("door_type", "")),
    }

    for field_name, (row_idx, col_idx) in fields.items():
        val = values.get(field_name, "")
        if field_name == "min_blank":
            continue  # leave empty
        cell, _ = _get_cell_at(table, row_idx, col_idx)
        if cell is not None:
            _set_cell_text(cell, val, font_pt=14)

    # Dropdowns: Tape (0), Inspect (1), SU Sheet (2)
    for i, key in enumerate(["tape", "inspect", "su_sheet"]):
        v = opts.get(key, "")
        if v:
            _set_sdt_dropdown_by_order(doc, i, v)

    # Machine checkboxes: Laser template = L1030,L95,L49,L3030S,L20; Combi = TC1000,TC3000
    if opts.get("is_laser", True):
        ticks = opts.get("laser_ticks", [False] * 5)
        for i, checked in enumerate(ticks[:5]):
            _set_sdt_checkbox_by_order(doc, i, checked)
    else:
        ticks = opts.get("combi_ticks", [False] * 2)
        for i, checked in enumerate(ticks[:2]):
            _set_sdt_checkbox_by_order(doc, i, checked)

    # NOTES field: fill with 1OFF PROGRAM or 2OFF PROGRAM when applicable (bold, 20pt)
    if notes_text:
        try:
            target, kind = _find_notes_content_target(doc)
            if target is not None:
                if kind == "tc":
                    _set_tc_text_bold(target, notes_text, font_pt=20)
                else:
                    _set_run_text_bold(target, notes_text, font_pt=20)
        except Exception:
            pass  # skip notes if template structure differs

    doc.save(output_path)
    log(f"  💾 {Path(output_path).name}")


def convert_to_pdf(docx_path, pdf_path, log):
    if HAS_PDF:
        # Suppress tqdm/progress output from docx2pdf (avoids terminal spam in exe)
        _devnull_out = open(os.devnull, "w")
        _devnull_err = open(os.devnull, "w")
        _old_stdout, _old_stderr = sys.stdout, sys.stderr
        try:
            sys.stdout, sys.stderr = _devnull_out, _devnull_err
            docx2pdf_convert(docx_path, pdf_path)
        finally:
            sys.stdout, sys.stderr = _old_stdout, _old_stderr
            _devnull_out.close()
            _devnull_err.close()
        log(f"  📄 {Path(pdf_path).name}")
    else:
        log("  ⚠️  docx2pdf not installed — .docx only")


def detect_template_type(docx_path: str) -> bool:
    """Detect if docx is Laser (True) or Combi (False) from document title."""
    if not HAS_DOCX:
        return True
    try:
        doc = Document(docx_path)
        for para in doc.paragraphs[:5]:
            t = (para.text or "").upper()
            if "COMBI" in t or "PUNCH" in t:
                return False
            if "LASER" in t:
                return True
        for table in doc.tables[:1]:
            for row in table.rows[:3]:
                for cell in row.cells:
                    t = (cell.text or "").upper()
                    if "COMBI" in t or "PUNCH" in t:
                        return False
                    if "LASER" in t:
                        return True
    except Exception:
        pass
    return True  # default laser


def update_docx_fields(docx_path: str, updates: dict, is_laser: bool, log=None) -> bool:
    """
    Update specific fields in an existing setup sheet docx.
    updates: {"pts_prog": "2", "pts_sht": "1", "tape": "NEW", ...}
    """
    if not HAS_DOCX:
        return False
    try:
        doc = Document(docx_path)
        fields = LASER_FIELDS if is_laser else COMBI_FIELDS
        table = doc.tables[0]

        # Table cell fields
        cell_fields = set(fields.keys()) - {"min_blank"}
        for fname, val in updates.items():
            if fname not in cell_fields:
                continue
            pos = fields.get(fname)
            if not pos:
                continue
            cell, _ = _get_cell_at(table, pos[0], pos[1])
            if cell is not None:
                _set_cell_text(cell, str(val), font_pt=14)
                if log:
                    log(f"  ✓ {fname} → {val}")

        # Dropdowns: tape, inspect, su_sheet
        dd_map = {"tape": 0, "inspect": 1, "su_sheet": 2}
        for fname, val in updates.items():
            if fname in dd_map and val:
                _set_sdt_dropdown_by_order(doc, dd_map[fname], str(val))
                if log:
                    log(f"  ✓ {fname} → {val}")

        # Checkboxes: laser_ticks (0-4) or combi_ticks (0-1)
        if "laser_ticks" in updates:
            ticks = updates["laser_ticks"]
            if isinstance(ticks, (list, tuple)) and is_laser:
                for i, chk in enumerate(ticks[:5]):
                    _set_sdt_checkbox_by_order(doc, i, bool(chk))
                if log:
                    log("  ✓ laser_ticks updated")
        if "combi_ticks" in updates:
            ticks = updates["combi_ticks"]
            if isinstance(ticks, (list, tuple)) and not is_laser:
                for i, chk in enumerate(ticks[:2]):
                    _set_sdt_checkbox_by_order(doc, i, bool(chk))
                if log:
                    log("  ✓ combi_ticks updated")

        # NOTES field: 1OFF PROGRAM / 2OFF PROGRAM (bold, 20pt)
        if "notes" in updates:
            notes_val = str(updates["notes"]).strip()
            if notes_val:
                target, kind = _find_notes_content_target(doc)
                if target is not None:
                    if kind == "tc":
                        _set_tc_text_bold(target, notes_val, font_pt=20)
                    else:
                        _set_run_text_bold(target, notes_val, font_pt=20)
                    if log:
                        log(f"  ✓ notes → {notes_val}")

        doc.save(docx_path)
        return True
    except Exception as e:
        if log:
            log(f"  ❌ Update error: {e}")
        return False


def ai_build_part_from_order(zz: str, chunk: str, is_door: bool, cfg: dict) -> str | None:
    """
    Ask AI to build part description from order text when regex rules fail.
    Returns e.g. ZZ9971952PXX_690x1887_S_F_RH_INTU or None on error.
    """
    part_type = "door (LASER)" if is_door else "frame (PUNCH)"
    custom_rules = cfg.get("verify_order_rules", [])
    rules_note = ""
    if custom_rules:
        rules_note = f"\nUser custom rules (pattern|replacement): {custom_rules}\n"
    prompt = f"""Parse this order line and build the Part No / part code.
{rules_note}

ZZ base: {zz}
Order text (from PDF):
---
{chunk[:1500]}
---

This is a {part_type}. Build the full part code in this format:
- Doors: ZZ_base_DIMxDIM_T1_LOCK_HAND (e.g. ZZ9971784PXX_774x1921_T1_EU_RH)
  LOCK: EU (Euro), STD (Standard), CL (Concealed)
  HAND: RH or LH
- Frames: ZZ_base_DIMxDIM_S_F_HAND_EXTRA or ZZ_base_DIMxDIM_D_F_HAND_EXTRA
  S_F=Single frame, D_F=Double frame
  HAND: RH, LH, or DR
  EXTRA: INTU for Intuframe (put last, e.g. S_F_RH_INTU)

Rules: TYPE N I LH = Single Frame Intuframe Left Hand. TYPENILH = same. Eurolock→EU. Dimensions like 690x1887.

Return ONLY the part code, nothing else. No explanation. Example: ZZ9971952PXX_690x1887_S_F_RH_INTU"""
    try:
        raw = ai_call(build_system_prompt(cfg), prompt, cfg, max_tokens=200)
        raw = (raw or "").strip()
        if not raw:
            return None
        raw = raw.split("\n")[0].strip()
        if raw.startswith("`"): raw = raw.strip("`")
        if zz.upper() in raw.upper() and re.search(r"\d{3,4}[xX]\d{3,4}", raw):
            return raw
        return None
    except Exception:
        return None


def ai_extract_field_changes(user_request: str, docx_path: str, cfg: dict) -> dict:
    """
    Ask AI to extract field name → value from user's fix request.
    Returns {"pts_prog": "2", "pts_sht": "1", ...}
    """
    prompt = f"""The user wants to fix a setup sheet document. They said:

"{user_request}"

Document: {Path(docx_path).name}

Editable fields (use these exact keys):
- pts_prog (Parts per Program, Pts/Prog)
- pts_sht (Parts per Sheet, Pts/Sht)
- thickness, type (material), sheet_size
- customer, part_number, description, revision, dnc, programmer
- tape (NEW/MODIFIED/N/A), inspect (FULLY/PARTLY/N/A), su_sheet (OLD/NEW)
- notes (NOTES field at bottom: "1OFF PROGRAM" or "2OFF PROGRAM" for 1OFF/2OFF parts)

Return ONLY valid JSON with field keys and new values. Example: {{"pts_prog": "2"}}
Only include fields the user wants to change. If unclear, return {{}}."""
    raw = ai_call(build_system_prompt(cfg), prompt, cfg, max_tokens=300)
    raw = re.sub(r"^```json|^```|```$", "", raw, flags=re.MULTILINE).strip()
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return {}


# ══════════════════════════════════════════════════════════════════════════
#  GUI
# ══════════════════════════════════════════════════════════════════════════

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("SheetGen  © Pruthvi")
        self.geometry("1100x860")
        self.minsize(900, 700)
        self.configure(bg=BG)

        self.cfg = load_config()

        # chat conversation history (list of {role, content})
        self._chat_history: list = []

        self._build_ui()
        self._check_deps()

    def _check_deps(self):
        missing = []
        if not HAS_EXCEL: missing.append("openpyxl")
        if not HAS_DOCX:  missing.append("python-docx")
        if not HAS_PDF:   missing.append("docx2pdf  (optional)")
        if missing:
            self.log("⚠️  Missing packages — install with pip:")
            for p in missing: self.log(f"   pip install {p}")
            self.log("")

    # ── top-level layout ──────────────────────────────────────────────────
    def _build_ui(self):
        # header
        hdr = tk.Frame(self, bg=BG)
        hdr.pack(fill="x", padx=22, pady=(16, 0))
        tk.Label(hdr, text="⚙", font=("Courier", 24), fg=ACCENT, bg=BG).pack(side="left")
        tf = tk.Frame(hdr, bg=BG)
        tf.pack(side="left", padx=10)
        tk.Label(tf, text="SHEETGEN",
                 font=("Courier", 14, "bold"), fg=TEXT, bg=BG).pack(anchor="w")
        tk.Label(tf, text="AI-powered  •  Laser & Punch Press  •  © Pruthvi",
                 font=("Courier", 8), fg=MUTED, bg=BG).pack(anchor="w")
        tk.Button(hdr, text="About", command=self._show_about,
                  bg=PANEL2, fg=MUTED, relief="flat", bd=0,
                  font=("Courier", 9), cursor="hand2",
                  activebackground=BORDER, activeforeground=TEXT,
                  padx=10, pady=5).pack(side="right", padx=(0, 4))
        tk.Button(hdr, text="⚙  Settings", command=self._open_settings,
                  bg=PANEL2, fg=MUTED, relief="flat", bd=0,
                  font=("Courier", 9), cursor="hand2",
                  activebackground=BORDER, activeforeground=TEXT,
                  padx=10, pady=5).pack(side="right")

        # provider badge
        self._provider_badge = tk.Label(hdr, text="", font=("Courier", 8),
                                        fg=BG, bg=ACCENT, padx=6, pady=3)
        self._provider_badge.pack(side="right", padx=(0, 6))
        self._update_provider_badge()

        tk.Frame(self, bg=ACCENT, height=1).pack(fill="x", padx=22, pady=8)

        # notebook tabs
        style = ttk.Style()
        style.theme_use("default")
        style.configure("TNotebook",        background=BG,    borderwidth=0)
        style.configure("TNotebook.Tab",    background=PANEL2, foreground=MUTED,
                        font=("Courier", 9, "bold"), padding=(16, 6))
        style.map("TNotebook.Tab",
                  background=[("selected", PANEL)],
                  foreground=[("selected", TEXT)])

        self._nb = ttk.Notebook(self)
        self._nb.pack(fill="both", expand=True, padx=22, pady=(0, 16))

        # Tab 1 — Generator
        gen_tab = tk.Frame(self._nb, bg=BG)
        self._nb.add(gen_tab, text="  ⚙  Generator  ")
        self._build_generator_tab(gen_tab)

        # Tab 2 — AI Chat
        chat_tab = tk.Frame(self._nb, bg=BG)
        self._nb.add(chat_tab, text="  💬  AI Assistant  ")
        self._build_chat_tab(chat_tab)

        # Tab 3 — Rename tool
        rename_tab = tk.Frame(self._nb, bg=BG)
        self._nb.add(rename_tab, text="  📁  Rename Tool  ")
        self._build_rename_tab(rename_tab)

        # Tab 4 — Verify Orders
        verify_tab = tk.Frame(self._nb, bg=BG)
        self._nb.add(verify_tab, text="  ✓  Verify Orders  ")
        self._build_verify_tab(verify_tab)

    # ══════════════════════════════════════════════════════════════════════
    #  TAB 1 — GENERATOR
    # ══════════════════════════════════════════════════════════════════════
    def _build_generator_tab(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(0, weight=1)

        # Main scrollable area
        main = tk.Frame(parent, bg=BG)
        main.grid(row=0, column=0, sticky="nsew", padx=16, pady=12)
        main.columnconfigure(0, weight=1)

        # Two columns: settings (left) + machines + log (right)
        content = tk.Frame(main, bg=BG)
        content.grid(row=0, column=0, sticky="nsew")
        content.columnconfigure(0, weight=1)
        content.columnconfigure(1, weight=0)
        content.rowconfigure(0, weight=1)

        left = tk.Frame(content, bg=BG)
        left.grid(row=0, column=0, sticky="nsew", padx=(0, 16))
        left.columnconfigure(0, weight=1)

        right = tk.Frame(content, bg=BG)
        right.grid(row=0, column=1, sticky="nsew")
        right.columnconfigure(0, weight=1)

        self._build_file_pickers(left)
        self._build_search_section(left)
        self._build_defaults_panel(left)
        self._build_dropdowns_panel(left)
        self._build_options(left)
        self._build_generate_btn(left)
        self._build_laser_ticks_panel(right)
        self._build_log_area(right)  # Activity log on right, under Laser/Combi
        self._load_profile()  # Apply active profile to UI

    def _build_file_pickers(self, parent):
        fp = tk.LabelFrame(parent, text=" Files ", font=("Courier", 8, "bold"),
            fg=ACCENT, bg=BG, bd=1, highlightbackground=BORDER)
        fp.grid(row=0, column=0, sticky="ew", pady=(0, 8))
        fp.columnconfigure(1, weight=1)

        # Profile row
        prof_row = tk.Frame(fp, bg=BG)
        prof_row.grid(row=0, column=0, columnspan=3, sticky="ew", padx=8, pady=(6, 4))
        prof_row.columnconfigure(1, weight=1)
        tk.Label(prof_row, text="Profile", font=("Courier", 8), fg=MUTED, bg=BG, width=10).pack(side="left", padx=(0, 6))
        self.profile_var = tk.StringVar(value=self.cfg.get("active_profile", "Default"))
        self._profile_dd = ttk.Combobox(prof_row, textvariable=self.profile_var, state="readonly",
                                        font=("Courier", 8), width=14)
        self._profile_dd.pack(side="left", fill="x", expand=True, padx=(0, 6))
        self._profile_dd.bind("<<ComboboxSelected>>", lambda e: self._load_profile())
        for btn_text, cmd in [("Save", self._save_profile), ("New", self._new_profile), ("Delete", self._delete_profile)]:
            tk.Button(prof_row, text=btn_text, command=cmd, bg=BORDER, fg=ACCENT, relief="flat",
                      font=("Courier", 7), cursor="hand2", padx=6, pady=2).pack(side="left", padx=2)
        self._refresh_profile_dd()

        self.excel_var       = tk.StringVar(value=self.cfg.get("last_excel", ""))
        self.laser_tpl_var   = tk.StringVar(value=self.cfg.get("last_laser_template", ""))
        self.combi_tpl_var   = tk.StringVar(value=self.cfg.get("last_combi_template", ""))
        self.output_var      = tk.StringVar(value=self.cfg.get("last_output",
                                             str(Path.home() / "Desktop")))
        self.start_row_var   = tk.IntVar(value=self.cfg.get("start_row", 1))
        self.max_rows_var    = tk.IntVar(value=self.cfg.get("max_rows", 15))

        def picker(label, var, filetypes, r):
            tk.Label(fp, text=label, font=("Courier", 8), fg=MUTED, bg=BG,
                    anchor="w", width=14).grid(row=r, column=0, sticky="w", padx=8, pady=3)
            tk.Entry(fp, textvariable=var, bg=PANEL2, fg=TEXT,
                     insertbackground=ACCENT, relief="flat", font=("Courier", 8),
                     bd=0, highlightthickness=1, highlightbackground=BORDER,
                     highlightcolor=ACCENT).grid(row=r, column=1, sticky="ew", padx=(0,6), pady=3, ipady=4)
            tk.Button(fp, text="…", command=lambda v=var, ft=filetypes: self._browse(v, ft),
                     bg=BORDER, fg=ACCENT, relief="flat", bd=0, font=("Courier", 8),
                     cursor="hand2", padx=6, pady=2).grid(row=r, column=2, pady=3)

        picker("Excel (Register)", self.excel_var, [("Excel","*.xlsx *.xls")], 1)
        picker("Laser template", self.laser_tpl_var, [("Word","*.docx")], 2)
        picker("Combi template", self.combi_tpl_var, [("Word","*.docx")], 3)
        picker("Output folder", self.output_var, None, 4)

        row_frame = tk.Frame(fp, bg=BG)
        row_frame.grid(row=5, column=0, columnspan=3, sticky="ew", padx=8, pady=(2, 6))
        row_frame.columnconfigure(1, weight=1)
        tk.Label(row_frame, text="Rows:", font=("Courier", 8), fg=MUTED, bg=BG).pack(side="left", padx=(0, 6))
        tk.Spinbox(row_frame, from_=1, to=1000, textvariable=self.start_row_var,
                   width=5, font=("Courier", 8), bg=PANEL2, fg=TEXT).pack(side="left", padx=(0, 4))
        tk.Label(row_frame, text="to", font=("Courier", 8), fg=MUTED, bg=BG).pack(side="left", padx=2)
        tk.Spinbox(row_frame, from_=1, to=500, textvariable=self.max_rows_var,
                   width=5, font=("Courier", 8), bg=PANEL2, fg=TEXT).pack(side="left", padx=2)
        tk.Label(row_frame, text="Filter:", font=("Courier", 8), fg=MUTED, bg=BG).pack(side="left", padx=(12, 4))
        self.filter_var = tk.StringVar(value=self.cfg.get("excel_filter", "All"))
        filter_dd = ttk.Combobox(row_frame, textvariable=self.filter_var, values=["All", "LASER only", "PUNCH only"],
                                 state="readonly", font=("Courier", 8), width=10)
        filter_dd.pack(side="left", padx=2)
        if self.filter_var.get() not in ["All", "LASER only", "PUNCH only"]:
            self.filter_var.set("All")
        tk.Label(row_frame, text="Max scan:", font=("Courier", 8), fg=MUTED, bg=BG).pack(side="left", padx=(8, 4))
        self.max_scan_var = tk.IntVar(value=self.cfg.get("max_scan", 500))
        tk.Spinbox(row_frame, from_=50, to=5000, textvariable=self.max_scan_var,
                   width=5, font=("Courier", 8), bg=PANEL2, fg=TEXT).pack(side="left", padx=2)

    def _build_search_section(self, parent):
        """Search by DNC across Laser + Punch; select individually for generation."""
        sf = tk.LabelFrame(parent, text=" Search by DNC (Laser + Punch) ", font=("Courier", 8, "bold"),
            fg=YELLOW, bg=BG, bd=1, highlightbackground=BORDER)
        sf.grid(row=1, column=0, sticky="ew", pady=(0, 8))
        sf.columnconfigure(1, weight=1)
        self.search_dnc_var = tk.StringVar(value="")
        self.search_results_data = []  # list of row dicts
        sr = tk.Frame(sf, bg=BG)
        sr.grid(row=0, column=0, columnspan=3, sticky="ew", padx=8, pady=(6, 4))
        sr.columnconfigure(1, weight=1)
        tk.Label(sr, text="DNC / Program No", font=("Courier", 8), fg=MUTED, bg=BG, width=14).pack(side="left", padx=(0, 6))
        tk.Entry(sr, textvariable=self.search_dnc_var, bg=PANEL2, fg=TEXT, font=("Courier", 8),
                 insertbackground=ACCENT, relief="flat", width=12).pack(side="left", fill="x", expand=True, padx=(0, 6), ipady=3)
        tk.Button(sr, text="Search", command=self._do_search, bg=ACCENT, fg=BG, relief="flat",
                  font=("Courier", 8), cursor="hand2", padx=8, pady=2).pack(side="left")
        list_frame = tk.Frame(sf, bg=BG)
        list_frame.grid(row=1, column=0, columnspan=3, sticky="nsew", padx=8, pady=(0, 6))
        list_frame.columnconfigure(0, weight=1)
        list_frame.rowconfigure(0, weight=1)
        self.search_listbox = tk.Listbox(list_frame, selectmode=tk.EXTENDED, height=6,
            bg=PANEL2, fg=TEXT, font=("Courier", 8), selectbackground=ACCENT, selectforeground=BG,
            highlightthickness=0, bd=0)
        sb = tk.Scrollbar(list_frame, orient="vertical", command=self.search_listbox.yview)
        self.search_listbox.configure(yscrollcommand=sb.set)
        self.search_listbox.grid(row=0, column=0, sticky="nsew")
        sb.grid(row=0, column=1, sticky="ns")
        btn_row = tk.Frame(sf, bg=BG)
        btn_row.grid(row=2, column=0, columnspan=3, sticky="ew", padx=8, pady=(0, 6))
        tk.Button(btn_row, text="Select all", command=self._search_select_all, bg=BORDER, fg=TEXT,
                  relief="flat", font=("Courier", 7), cursor="hand2", padx=6, pady=2).pack(side="left", padx=(0, 4))
        tk.Button(btn_row, text="Clear", command=self._search_clear, bg=BORDER, fg=TEXT,
                  relief="flat", font=("Courier", 7), cursor="hand2", padx=6, pady=2).pack(side="left")
        tk.Label(sf, text="Select rows above to generate those only. Or use Filter below.", font=("Courier", 7),
                 fg=MUTED, bg=BG).grid(row=3, column=0, columnspan=3, sticky="w", padx=8, pady=(0, 6))

    def _do_search(self):
        ep = self.excel_var.get().strip()
        q = self.search_dnc_var.get().strip()
        if not ep or not os.path.exists(ep):
            messagebox.showerror("Missing", "Select a valid Excel file first.")
            return
        if not q:
            messagebox.showinfo("Search", "Enter a DNC/Program number to search.")
            return
        self.log("\n" + "─"*40)
        self.log(f"🔍 Searching for '{q}' in Laser + Punch…")
        try:
            # Search uses high limit (100k) to find DNC anywhere; UI Max scan is for filter mode only
            rows = search_register_by_dnc(ep, q, self.log, max_scan=100000)
            self.search_results_data = rows
            self.search_listbox.delete(0, tk.END)
            for r in rows:
                dnc = r.get("dnc", "")
                lp = r.get("laser_punch", "")
                pc = (r.get("part_code", "") or "")[:45]
                if len((r.get("part_code", "") or "")) > 45:
                    pc += "…"
                self.search_listbox.insert(tk.END, f"{dnc} | {lp} | {pc}")
            if rows:
                self.log(f"  ✅ {len(rows)} result(s) — select which to generate, or use Filter mode.")
            else:
                self.log("  No matches found.")
        except Exception as e:
            self.log(f"  ❌ Search error: {e}")
            messagebox.showerror("Search error", str(e))

    def _search_select_all(self):
        self.search_listbox.selection_set(0, tk.END)

    def _search_clear(self):
        self.search_listbox.selection_clear(0, tk.END)

    def _get_selected_search_rows(self):
        """Return list of row dicts for currently selected search results."""
        sel = self.search_listbox.curselection()
        if not sel:
            return []
        return [self.search_results_data[i] for i in sel if 0 <= i < len(self.search_results_data)]

    def _refresh_profile_dd(self):
        profiles = list(self.cfg.get("profiles", {}).keys())
        self._profile_dd["values"] = profiles
        if profiles and self.profile_var.get() not in profiles:
            self.profile_var.set(profiles[0])

    def _safe_max_scan(self):
        try:
            return max(50, min(5000, int(self.max_scan_var.get() or 500)))
        except (TypeError, ValueError, tk.TclError):
            return 500

    def _get_current_ui_data(self):
        """Capture current UI state as profile data."""
        return {
            "programmer": self.programmer_var.get(),
            "sheet_size": self.sheet_size_var.get(),
            "pts_prog": self.pts_prog_var.get(),
            "pts_sht": self.pts_sht_var.get(),
            "thickness": self.thickness_var.get(),
            "material_type": self.material_type_var.get(),
            "tape": self.tape_var.get() or "NEW",
            "inspect": self.inspect_var.get() or "FULLY",
            "su_sheet": self.su_sheet_var.get() or "NEW",
            "last_excel": self.excel_var.get(),
            "last_laser_template": self.laser_tpl_var.get(),
            "last_combi_template": self.combi_tpl_var.get(),
            "last_output": self.output_var.get(),
            "start_row": self.start_row_var.get(),
            "max_rows": self.max_rows_var.get(),
            "excel_filter": (self.filter_var.get() or "All").strip(),
            "max_scan": self._safe_max_scan(),
            "laser_ticks": [v.get() for v in self.laser_ticks],
            "combi_ticks": [v.get() for v in getattr(self, "combi_ticks", [])],
        }

    def _apply_profile_data(self, data: dict):
        """Apply profile data to UI."""
        self.programmer_var.set(data.get("programmer", ""))
        self.sheet_size_var.set(data.get("sheet_size", "2500x1250"))
        self.pts_prog_var.set(data.get("pts_prog", ""))
        self.pts_sht_var.set(data.get("pts_sht", ""))
        self.thickness_var.set(data.get("thickness", "1.2"))
        self.material_type_var.set(data.get("material_type", "Galv"))
        self.tape_var.set(data.get("tape", "NEW"))
        self.inspect_var.set(data.get("inspect", "FULLY"))
        self.su_sheet_var.set(data.get("su_sheet", "NEW"))
        self.excel_var.set(data.get("last_excel", ""))
        self.laser_tpl_var.set(data.get("last_laser_template", ""))
        self.combi_tpl_var.set(data.get("last_combi_template", ""))
        self.output_var.set(data.get("last_output", "") or str(Path.home() / "Desktop"))
        self.start_row_var.set(max(1, data.get("start_row", 1)))
        self.max_rows_var.set(max(1, data.get("max_rows", 15)))
        fv = (data.get("excel_filter") or "All").strip()
        if fv in ["All", "LASER only", "PUNCH only"]:
            self.filter_var.set(fv)
        else:
            self.filter_var.set("All")
        self.max_scan_var.set(max(50, min(5000, data.get("max_scan", 500))))
        for i, v in enumerate((data.get("laser_ticks") or [False] * 5)[:5]):
            if i < len(self.laser_ticks):
                self.laser_ticks[i].set(v)
        combi = getattr(self, "combi_ticks", [])
        for i, v in enumerate((data.get("combi_ticks") or [False] * 2)[:2]):
            if i < len(combi):
                combi[i].set(v)

    def _load_profile(self):
        name = self.profile_var.get()
        profiles = self.cfg.get("profiles", {})
        if name in profiles:
            self.cfg["active_profile"] = name
            self._apply_profile_data(profiles[name])
            save_config(self.cfg)

    def _save_profile(self):
        name = self.profile_var.get()
        if not name:
            return
        profiles = self.cfg.get("profiles", {})
        profiles[name] = self._get_current_ui_data()
        self.cfg["profiles"] = profiles
        self.cfg["active_profile"] = name
        save_config(self.cfg)
        self._refresh_profile_dd()
        messagebox.showinfo("Profile saved", f"Saved to profile \"{name}\".")

    def _new_profile(self):
        name = simpledialog.askstring("New profile", "Profile name (e.g. ASPEX Laser, Caterpillar, NiftyLifts):")
        if not name or not name.strip():
            return
        name = name.strip()
        profiles = self.cfg.get("profiles", {})
        if name in profiles:
            if not messagebox.askyesno("Overwrite?", f"Profile \"{name}\" exists. Overwrite?"):
                return
        profiles[name] = self._get_current_ui_data()
        self.cfg["profiles"] = profiles
        self.cfg["active_profile"] = name
        self.profile_var.set(name)
        save_config(self.cfg)
        self._refresh_profile_dd()
        messagebox.showinfo("Profile created", f"Created profile \"{name}\".")

    def _delete_profile(self):
        name = self.profile_var.get()
        profiles = self.cfg.get("profiles", {})
        if name not in profiles or len(profiles) <= 1:
            messagebox.showinfo("Cannot delete", "Cannot delete the only profile.")
            return
        if not messagebox.askyesno("Delete profile", f"Delete profile \"{name}\"?"):
            return
        del profiles[name]
        self.cfg["profiles"] = profiles
        self.cfg["active_profile"] = list(profiles.keys())[0]
        self.profile_var.set(self.cfg["active_profile"])
        save_config(self.cfg)
        self._refresh_profile_dd()
        self._load_profile()
        messagebox.showinfo("Deleted", f"Profile \"{name}\" deleted.")

    def _build_defaults_panel(self, parent):
        df = tk.LabelFrame(parent, text=" Sheet defaults ", font=("Courier", 8, "bold"),
            fg=ACCENT, bg=BG, bd=1, highlightbackground=BORDER)
        df.grid(row=2, column=0, sticky="ew", pady=(0, 8))
        df.columnconfigure(1, weight=1)
        df.columnconfigure(3, weight=1)
        cfg = self.cfg
        self.programmer_var   = tk.StringVar(value=cfg.get("programmer", ""))
        self.sheet_size_var   = tk.StringVar(value=cfg.get("sheet_size", "2500x1250"))
        self.pts_prog_var     = tk.StringVar(value=cfg.get("pts_prog", ""))
        self.pts_sht_var      = tk.StringVar(value=cfg.get("pts_sht", ""))
        self.thickness_var    = tk.StringVar(value=cfg.get("thickness", "1.2"))
        self.material_type_var= tk.StringVar(value=cfg.get("material_type", "Galv"))
        pairs = [
            ("Programmer", self.programmer_var), ("Sheet Size", self.sheet_size_var),
            ("Pts/Prog", self.pts_prog_var), ("Pts/Sht", self.pts_sht_var),
            ("Thickness", self.thickness_var), ("Type", self.material_type_var),
        ]
        for i, (label, var) in enumerate(pairs):
            r, c = i // 2, (i % 2) * 2
            tk.Label(df, text=label, font=("Courier", 8), fg=MUTED, bg=BG, width=10).grid(
                row=r, column=c, sticky="w", padx=(8 if c == 0 else 16, 4), pady=2)
            tk.Entry(df, textvariable=var, bg=PANEL2, fg=TEXT, font=("Courier", 8),
                     width=10).grid(row=r, column=c+1, sticky="ew", padx=(0, 8 if c == 0 else 8), pady=2)

    def _build_dropdowns_panel(self, parent):
        dd = tk.LabelFrame(parent, text=" Tape · Inspect · SU Sheet ", font=("Courier", 8, "bold"),
            fg=YELLOW, bg=BG, bd=1, highlightbackground=BORDER)
        dd.grid(row=3, column=0, sticky="ew", pady=(0, 8))
        dd.columnconfigure(1, weight=1)
        dd.columnconfigure(3, weight=1)
        dd.columnconfigure(5, weight=1)
        self.tape_var    = tk.StringVar(value=self.cfg.get("tape", "NEW"))
        self.inspect_var = tk.StringVar(value=self.cfg.get("inspect", "FULLY"))
        self.su_sheet_var= tk.StringVar(value=self.cfg.get("su_sheet", "NEW"))
        opts_tape = ["Choose an item.", "NEW", "MODIFIED", "N/A"]
        opts_inspect = ["Choose an item.", "FULLY", "PARTLY", "N/A"]
        opts_su = ["Choose an item.", "OLD", "NEW"]
        for c, (label, var, opts) in enumerate([
            ("Tape", self.tape_var, opts_tape),
            ("Inspect", self.inspect_var, opts_inspect),
            ("SU Sheet", self.su_sheet_var, opts_su)]):
            tk.Label(dd, text=label, font=("Courier", 8), fg=MUTED, bg=BG).grid(
                row=0, column=c*2, sticky="w", padx=(8 if c == 0 else 12, 4), pady=6)
            om = ttk.Combobox(dd, textvariable=var, values=opts, state="readonly",
                              font=("Courier", 8), width=10)
            om.grid(row=0, column=c*2+1, sticky="ew", padx=(0, 8), pady=6)

    def _build_laser_ticks_panel(self, parent):
        """Machine selection: Laser (5) + Combi/Punch (2)."""
        container = tk.Frame(parent, bg=BG)
        container.grid(row=0, column=0, sticky="nw")
        parent.columnconfigure(0, weight=1)

        laser_frame = tk.LabelFrame(container, text=" Laser (when Excel says LASER) ",
            font=("Courier", 8, "bold"), fg=ACCENT, bg=BG, bd=1, highlightbackground=BORDER)
        laser_frame.pack(fill="x", pady=(0, 6))
        laser_frame.columnconfigure(0, weight=1)
        self.laser_ticks = []
        for i, label in enumerate(["L1030(L22)", "L95", "L49", "L3030S", "L20"]):
            v = tk.BooleanVar(value=False)
            self.laser_ticks.append(v)
            r, c = i // 3, i % 3
            tk.Checkbutton(laser_frame, text=label, variable=v,
                bg=BG, fg=TEXT, selectcolor=PANEL2, activebackground=BG, activeforeground=ACCENT,
                font=("Courier", 8), cursor="hand2").grid(row=r, column=c, sticky="w", padx=8, pady=3)

        combi_frame = tk.LabelFrame(container, text=" Combi (when Excel says PUNCH) ",
            font=("Courier", 8, "bold"), fg=YELLOW, bg=BG, bd=1, highlightbackground=BORDER)
        combi_frame.pack(fill="x", pady=(0, 4))
        self.combi_ticks = []
        for i, label in enumerate(["TC1000", "TC3000"]):
            v = tk.BooleanVar(value=False)
            self.combi_ticks.append(v)
            tk.Checkbutton(combi_frame, text=label, variable=v,
                bg=BG, fg=TEXT, selectcolor=PANEL2, activebackground=BG, activeforeground=ACCENT,
                font=("Courier", 8), cursor="hand2").grid(row=0, column=i, sticky="w", padx=8, pady=6)

    def _build_options(self, parent):
        opt = tk.Frame(parent, bg=BG)
        opt.grid(row=4, column=0, sticky="ew", pady=(0, 8))
        self.save_pdf_var     = tk.BooleanVar(value=True)
        self.save_docx_var    = tk.BooleanVar(value=True)
        self.use_ai_fallback_var = tk.BooleanVar(value=self.cfg.get("use_ai_fallback", True))
        for text, var in [("Save PDF", self.save_pdf_var),
                          ("Save .docx", self.save_docx_var),
                          ("Use AI when struggling", self.use_ai_fallback_var)]:
            tk.Checkbutton(opt, text=text, variable=var,
                           bg=BG, fg=TEXT, selectcolor=PANEL2,
                           activebackground=BG, activeforeground=ACCENT,
                           font=("Courier", 8), cursor="hand2").pack(side="left", padx=(0, 12))

    def _build_generate_btn(self, parent):
        bf = tk.Frame(parent, bg=BG)
        bf.grid(row=5, column=0, sticky="ew", pady=(0, 12))
        self.gen_btn = tk.Button(bf,
            text="▶   GENERATE SETUP SHEETS",
            command=self._run,
            bg=ACCENT, fg=BG, relief="flat", bd=0,
            font=("Courier", 11, "bold"), cursor="hand2",
            activebackground="#00b8cc", activeforeground=BG, pady=11)
        self.gen_btn.pack(fill="x")
        self.progress = ttk.Progressbar(bf, mode="indeterminate")
        self.progress.pack(fill="x", pady=(4, 0))
        s = ttk.Style()
        s.configure("TProgressbar", troughcolor=PANEL, background=ACCENT, thickness=3)

    def _build_log_area(self, parent):
        """Activity log — placed under Laser/Combi machine choices on the right."""
        parent.rowconfigure(2, weight=1)
        tk.Label(parent, text="Activity log",
                 font=("Courier", 8, "bold"), fg=MUTED, bg=BG).grid(
                     row=1, column=0, sticky="w", pady=(12, 4))
        self.log_box = scrolledtext.ScrolledText(
            parent, bg=PANEL, fg="#a0ffb0",
            font=("Courier", 8), relief="flat",
            insertbackground=ACCENT, bd=0,
            highlightthickness=1, highlightbackground=BORDER, state="disabled",
            width=36)  # Narrower so left panel isn't squashed
        self.log_box.grid(row=2, column=0, sticky="nsew", pady=(0, 8))
        self.log_box.tag_config("err",  foreground=RED)
        self.log_box.tag_config("ok",   foreground=ACCENT)
        self.log_box.tag_config("warn", foreground=YELLOW)

    # ══════════════════════════════════════════════════════════════════════
    #  TAB 2 — AI CHAT
    # ══════════════════════════════════════════════════════════════════════
    def _build_chat_tab(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(1, weight=1)

        # ── top: learned rules banner ──
        self._rules_frame = tk.Frame(parent, bg=PANEL2)
        self._rules_frame.grid(row=0, column=0, sticky="ew", padx=8, pady=(8,0))
        self._build_rules_banner()

        # ── middle: chat display ──
        chat_outer = tk.Frame(parent, bg=BG)
        chat_outer.grid(row=1, column=0, sticky="nsew", padx=8, pady=6)
        chat_outer.rowconfigure(0, weight=1)
        chat_outer.columnconfigure(0, weight=1)

        self._chat_canvas = tk.Canvas(chat_outer, bg=BG, bd=0, highlightthickness=0)
        chat_sb = tk.Scrollbar(chat_outer, orient="vertical",
                               command=self._chat_canvas.yview)
        self._chat_canvas.configure(yscrollcommand=chat_sb.set)
        chat_sb.grid(row=0, column=1, sticky="ns")
        self._chat_canvas.grid(row=0, column=0, sticky="nsew")

        self._chat_inner = tk.Frame(self._chat_canvas, bg=BG)
        self._chat_cwin  = self._chat_canvas.create_window(
            (0,0), window=self._chat_inner, anchor="nw")
        self._chat_inner.bind("<Configure>",
            lambda e: self._chat_canvas.configure(
                scrollregion=self._chat_canvas.bbox("all")))
        self._chat_canvas.bind("<Configure>",
            lambda e: self._chat_canvas.itemconfig(self._chat_cwin, width=e.width))
        self._chat_canvas.bind_all("<MouseWheel>",
            lambda e: self._chat_canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

        # welcome message
        self._add_chat_bubble(
            "assistant",
            "👋  Hi! I'm your setup sheet AI assistant.\n\n"
            "I already know the rules for doors (LASER) and frames (PUNCH/Combi).\n\n"
            "Fix sheets: \"P44302 thickness wrong, change to 2\" or \"Fix P44300, P44301, P44302 set pts/prog to 2\" "
            "or \"Change thickness to 2 for all sheets\". Uses Output folder from Generator. 🧠"
        )

        # ── bottom: input area ──
        input_area = tk.Frame(parent, bg=PANEL2)
        input_area.grid(row=2, column=0, sticky="ew", padx=8, pady=(0,8))
        input_area.columnconfigure(0, weight=1)

        # quick action buttons
        quick = tk.Frame(input_area, bg=PANEL2)
        quick.grid(row=0, column=0, columnspan=2, sticky="ew", padx=8, pady=(8,4))

        for label, msg in [
            ("🔍 Explain last parse",  "Can you explain how you parsed the last part code? Walk me through what each segment meant."),
            ("⚠️ Report a mistake",    "I want to report a mistake in how you parsed or filled the template."),
            ("➕ Add rule",            "SYSTEM_ADD_RULE"),
            ("📋 Show learned rules",  "What rules have you learned from my feedback so far?"),
            ("🔄 Reset to built-in",   "SYSTEM_RESET_BUILTIN"),
            ("🗑 Clear all rules",     "SYSTEM_CLEAR_RULES"),
        ]:
            tk.Button(quick, text=label,
                      command=lambda m=msg: self._quick_action(m),
                      bg=BORDER, fg=TEXT, relief="flat", bd=0,
                      font=("Courier", 8), cursor="hand2",
                      activebackground=PANEL2, activeforeground=ACCENT,
                      padx=8, pady=4).pack(side="left", padx=(0, 6))

        # text input
        self._chat_input = tk.Text(input_area, bg=PANEL, fg=TEXT,
                                   insertbackground=ACCENT, relief="flat",
                                   font=("Courier", 9), bd=0,
                                   highlightthickness=1,
                                   highlightbackground=BORDER,
                                   highlightcolor=ACCENT,
                                   height=3, wrap="word")
        self._chat_input.grid(row=1, column=0, sticky="ew",
                              padx=(8,6), pady=(0,8), ipady=6)
        self._chat_input.bind("<Return>",    self._on_enter)
        self._chat_input.bind("<Shift-Return>", lambda e: None)  # allow newline

        send_btn = tk.Button(input_area, text="Send\n↵",
                             command=self._send_chat,
                             bg=ACCENT, fg=BG, relief="flat", bd=0,
                             font=("Courier", 9, "bold"), cursor="hand2",
                             activebackground="#00b8cc", activeforeground=BG,
                             padx=12, pady=6)
        send_btn.grid(row=1, column=1, padx=(0,8), pady=(0,8), sticky="ns")

        tk.Label(input_area,
                 text="Enter = send  •  Shift+Enter = new line",
                 font=("Courier", 7), fg=MUTED, bg=PANEL2).grid(
                     row=2, column=0, columnspan=2, sticky="w", padx=8, pady=(0,4))

    # ══════════════════════════════════════════════════════════════════════
    #  TAB 3 — RENAME TOOL
    # ══════════════════════════════════════════════════════════════════════
    def _build_rename_tab(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(3, weight=1)

        # ── top: instructions ──
        inst = tk.Label(parent,
            text="Rename PDFs by DNC: extracts DNC from each PDF (text or OCR for scanned docs), looks up Part No in Excel, renames file. For scanned PDFs, install Tesseract and: pip install pymupdf pytesseract Pillow",
            font=("Courier", 8), fg=MUTED, bg=BG, wraplength=520)
        inst.grid(row=0, column=0, sticky="ew", padx=16, pady=(12, 4))

        # ── settings ──
        frm = tk.LabelFrame(parent, text=" Settings ", font=("Courier", 8, "bold"),
            fg=ACCENT, bg=BG, bd=1, highlightbackground=BORDER)
        frm.grid(row=1, column=0, sticky="ew", padx=16, pady=8)
        frm.columnconfigure(1, weight=1)

        self.rename_folder_var = tk.StringVar(value="")
        self.rename_excel_var = tk.StringVar(value=self.cfg.get("last_excel", ""))

        def row(lbl, var, r, is_folder=False):
            tk.Label(frm, text=lbl, font=("Courier", 8), fg=MUTED, bg=BG, width=12).grid(
                row=r, column=0, sticky="w", padx=8, pady=4)
            e = tk.Entry(frm, textvariable=var, bg=PANEL2, fg=TEXT, font=("Courier", 8),
                relief="flat", bd=0, highlightthickness=1, highlightbackground=BORDER)
            e.grid(row=r, column=1, sticky="ew", padx=(0, 6), pady=4, ipady=4)
            def pick():
                if is_folder:
                    p = filedialog.askdirectory(title="Select folder with PDFs")
                else:
                    p = filedialog.askopenfilename(
                        title="Select Register Excel",
                        filetypes=[("Excel", "*.xlsx *.xls"), ("All", "*.*")])
                if p:
                    if is_folder:
                        var.set(p)
                    else:
                        self._sync_excel_path(p)
            tk.Button(frm, text="Browse", command=pick, bg=BORDER, fg=ACCENT,
                relief="flat", font=("Courier", 7), cursor="hand2", padx=6, pady=2).grid(
                row=r, column=2, padx=(0, 8), pady=4)
            return e

        row("PDF folder", self.rename_folder_var, 0, is_folder=True)
        row("Register Excel", self.rename_excel_var, 1, is_folder=False)

        # Rename button
        btn_frm = tk.Frame(parent, bg=BG)
        btn_frm.grid(row=2, column=0, sticky="ew", padx=16, pady=(0, 8))
        tk.Button(btn_frm, text="  Rename PDFs  ", command=self._run_rename_tool,
            bg=ACCENT, fg=BG, relief="flat", font=("Courier", 9, "bold"),
            cursor="hand2", padx=16, pady=6,
            activebackground="#00b8cc", activeforeground=BG).pack(side="left")

        # Log
        log_lf = tk.LabelFrame(parent, text=" Log ", font=("Courier", 8, "bold"),
            fg=ACCENT, bg=BG, bd=1, highlightbackground=BORDER)
        log_lf.grid(row=3, column=0, sticky="nsew", padx=16, pady=(0, 12))
        log_lf.columnconfigure(0, weight=1)
        log_lf.rowconfigure(0, weight=1)
        self.rename_log = scrolledtext.ScrolledText(log_lf, bg=PANEL2, fg=TEXT,
            font=("Courier", 8), height=12, wrap="word", state="disabled",
            relief="flat", bd=0, highlightthickness=0)
        self.rename_log.grid(row=0, column=0, sticky="nsew", padx=6, pady=6)
        self.rename_log.tag_config("err", foreground=RED)
        self.rename_log.tag_config("ok", foreground=GREEN)
        self.rename_log.tag_config("warn", foreground=YELLOW)

    def _run_rename_tool(self):
        folder = self.rename_folder_var.get().strip()
        excel_path = self.rename_excel_var.get().strip()
        log = self.rename_log

        def append(msg, tag=None):
            log.config(state="normal")
            if tag:
                log.insert("end", msg + "\n", tag)
            else:
                log.insert("end", msg + "\n")
            log.see("end")
            log.config(state="disabled")

        if not folder or not os.path.isdir(folder):
            append("Please select a valid PDF folder.", "err")
            return
        if not excel_path or not os.path.isfile(excel_path):
            append("Please select a valid Register Excel file.", "err")
            return
        if not HAS_PYPDF:
            append("pypdf is required. Run: pip install pypdf", "err")
            return

        append("Building DNC → Part No lookup from Excel…")
        mapping = build_dnc_to_part_code(excel_path)
        append(f"  Loaded {len(mapping)} DNC(s) from Register.")

        pdfs = [f for f in os.listdir(folder) if f.lower().endswith(".pdf")]
        if not pdfs:
            append("No PDF files found in folder.", "warn")
            return

        append(f"Found {len(pdfs)} PDF(s). Processing…")
        ok, skip, fail = 0, 0, 0
        ocr_hint_shown = False
        for fn in sorted(pdfs):
            path = os.path.join(folder, fn)
            dnc = extract_dnc_from_pdf(path)
            if not dnc:
                hint = ""
                if not ocr_hint_shown and not (HAS_PYMUPDF and HAS_PYTESSERACT):
                    hint = " (for scanned PDFs: pip install pymupdf pytesseract Pillow, and install Tesseract)"
                    ocr_hint_shown = True
                append(f"  ⚠ {fn}: no DNC found in PDF{hint}", "warn")
                skip += 1
                continue
            part = mapping.get(dnc)
            if not part:
                append(f"  ⚠ {fn}: DNC {dnc} not in Register", "warn")
                skip += 1
                continue
            safe = sanitize_filename(part) + ".pdf"
            dest = os.path.join(folder, safe)
            if dest == path:
                append(f"  ✓ {fn}: already named correctly", "ok")
                ok += 1
                continue
            if os.path.exists(dest):
                append(f"  ⚠ {fn}: target {safe} already exists, skipping", "warn")
                skip += 1
                continue
            try:
                os.rename(path, dest)
                append(f"  ✓ {fn} → {safe}", "ok")
                ok += 1
            except Exception as e:
                append(f"  ✗ {fn}: {e}", "err")
                fail += 1
        append(f"Done: {ok} renamed, {skip} skipped, {fail} failed.")

    # ══════════════════════════════════════════════════════════════════════
    #  TAB 4 — VERIFY ORDERS
    # ══════════════════════════════════════════════════════════════════════
    def _build_verify_tab(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(1, weight=1)
        parent.rowconfigure(2, weight=0)

        # ── top: compact header + inputs ──
        top = tk.Frame(parent, bg=BG)
        top.grid(row=0, column=0, sticky="ew", padx=20, pady=(16, 12))
        top.columnconfigure(1, weight=1)

        tk.Label(top, text="Verify Order", font=("Courier", 11, "bold"),
                 fg=TEXT, bg=BG).grid(row=0, column=0, columnspan=3, sticky="w", pady=(0, 8))
        tk.Label(top, text="Compare order PDF ZZ numbers against Register (all sheets).",
                 font=("Courier", 8), fg=MUTED, bg=BG).grid(row=1, column=0, columnspan=3, sticky="w", pady=(0, 12))

        self.verify_pdf_var = tk.StringVar(value="")
        self.verify_excel_var = tk.StringVar(value=self.cfg.get("last_excel", ""))

        def file_row(lbl, var, r, is_pdf=False):
            tk.Label(top, text=lbl, font=("Courier", 8), fg=MUTED, bg=BG, width=14).grid(
                row=r, column=0, sticky="w", padx=(0, 10), pady=6)
            e = tk.Entry(top, textvariable=var, bg=PANEL2, fg=TEXT, font=("Courier", 8),
                relief="flat", bd=0, highlightthickness=1, highlightbackground=BORDER)
            e.grid(row=r, column=1, sticky="ew", padx=(0, 8), pady=6, ipady=5)
            def pick():
                if is_pdf:
                    p = filedialog.askopenfilename(title="Select Order PDF",
                        filetypes=[("PDF", "*.pdf"), ("All", "*.*")])
                else:
                    p = filedialog.askopenfilename(title="Select Register Excel",
                        filetypes=[("Excel", "*.xlsx *.xls"), ("All", "*.*")])
                if p:
                    if is_pdf:
                        var.set(p)
                    else:
                        self._sync_excel_path(p)
            tk.Button(top, text="Browse", command=pick, bg=BORDER, fg=ACCENT,
                relief="flat", font=("Courier", 8), cursor="hand2",
                padx=12, pady=4, activebackground=PANEL2, activeforeground=ACCENT).grid(
                row=r, column=2, padx=(0, 0), pady=6)

        file_row("Order PDF", self.verify_pdf_var, 2, is_pdf=True)
        file_row("Register Excel", self.verify_excel_var, 3, is_pdf=False)

        btn_frm = tk.Frame(top, bg=BG)
        btn_frm.grid(row=4, column=0, columnspan=3, sticky="w", pady=(0, 8))
        tk.Button(btn_frm, text="  Verify Order  ", command=self._run_verify_orders,
            bg=ACCENT, fg=BG, relief="flat", font=("Courier", 9, "bold"),
            cursor="hand2", padx=20, pady=8,
            activebackground="#00b8cc", activeforeground=BG).pack(side="left", padx=(0, 8))
        tk.Button(btn_frm, text="  🤖 AI Suggest  ", command=self._on_verify_ai_suggest,
            bg=PANEL2, fg=ACCENT, relief="flat", font=("Courier", 9, "bold"),
            cursor="hand2", padx=12, pady=8,
            activebackground=BORDER, activeforeground=ACCENT).pack(side="left", padx=(0, 8))
        tk.Button(btn_frm, text="  Add to Register  ", command=self._on_verify_add_selected,
            bg=BORDER, fg=ACCENT, relief="flat", font=("Courier", 9, "bold"),
            cursor="hand2", padx=16, pady=8,
            activebackground=PANEL2, activeforeground=ACCENT).pack(side="left")

        # ── results: two panels side by side ──
        results_frm = tk.Frame(parent, bg=BG)
        results_frm.grid(row=1, column=0, sticky="nsew", padx=20, pady=(0, 20))
        results_frm.columnconfigure(0, weight=1)
        results_frm.columnconfigure(1, weight=1)
        results_frm.rowconfigure(1, weight=1)

        # Summary bar
        self.verify_summary = tk.Label(results_frm, text="Select files and click Verify Order.",
            font=("Courier", 9), fg=MUTED, bg=BG)
        self.verify_summary.grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 8))

        # NEW panel
        new_lf = tk.LabelFrame(results_frm, text=" New (not in Register) — double-click or Add to Register ", font=("Courier", 9, "bold"),
            fg=YELLOW, bg=BG, bd=1, highlightbackground=BORDER)
        new_lf.grid(row=1, column=0, sticky="nsew", padx=(0, 8), pady=0)
        new_lf.columnconfigure(0, weight=1)
        new_lf.rowconfigure(0, weight=1)
        self.verify_new_list = tk.Listbox(new_lf, bg=PANEL2, fg=YELLOW, font=("Courier", 9),
            selectmode="browse", relief="flat", bd=0, highlightthickness=0)
        self.verify_new_list.grid(row=0, column=0, sticky="nsew", padx=6, pady=6)
        self.verify_new_list.bind("<Double-Button-1>", lambda e: self._on_verify_add_selected())
        new_sb = tk.Scrollbar(new_lf, orient="vertical", command=self.verify_new_list.yview)
        new_sb.grid(row=0, column=1, sticky="ns")
        self.verify_new_list.configure(yscrollcommand=new_sb.set)

        # OLD panel
        old_lf = tk.LabelFrame(results_frm, text=" Existing (in Register) ", font=("Courier", 9, "bold"),
            fg=GREEN, bg=BG, bd=1, highlightbackground=BORDER)
        old_lf.grid(row=1, column=1, sticky="nsew", padx=(8, 0), pady=0)
        old_lf.columnconfigure(0, weight=1)
        old_lf.rowconfigure(0, weight=1)
        self.verify_old_list = tk.Listbox(old_lf, bg=PANEL2, fg=GREEN, font=("Courier", 9),
            selectmode="extended", relief="flat", bd=0, highlightthickness=0)
        self.verify_old_list.grid(row=0, column=0, sticky="nsew", padx=6, pady=6)
        old_sb = tk.Scrollbar(old_lf, orient="vertical", command=self.verify_old_list.yview)
        old_sb.grid(row=0, column=1, sticky="ns")
        self.verify_old_list.configure(yscrollcommand=old_sb.set)

        # ── Verify Order Rules ──
        rules_lf = tk.LabelFrame(parent, text=" Verify Order Rules (add custom rules) ", font=("Courier", 9, "bold"),
            fg=ACCENT, bg=BG, bd=1, highlightbackground=BORDER)
        rules_lf.grid(row=2, column=0, sticky="ew", padx=20, pady=(0, 16))
        rules_lf.columnconfigure(0, weight=1)
        tk.Label(rules_lf,
            text="Format: pattern|replacement (e.g. TYPE N|S_F or Intuframe|INTU). Built-in: A1/A2=T1, Euro=EU, TYPE N I LH=S_F+INTU. Double-click to remove.",
            font=("Courier", 7), fg=MUTED, bg=BG, wraplength=600).grid(row=0, column=0, sticky="w", padx=8, pady=(6, 2))
        rules_inner = tk.Frame(rules_lf, bg=BG)
        rules_inner.grid(row=1, column=0, sticky="ew", padx=8, pady=(0, 6))
        rules_inner.columnconfigure(0, weight=1)
        self.verify_rules_var = tk.StringVar(value="")
        tk.Entry(rules_inner, textvariable=self.verify_rules_var, bg=PANEL2, fg=TEXT, font=("Courier", 8),
            relief="flat", bd=0, highlightthickness=1, highlightbackground=BORDER).grid(
            row=0, column=0, sticky="ew", padx=(0, 8), ipady=4)
        tk.Button(rules_inner, text="Add Rule", command=self._add_verify_rule,
            bg=BORDER, fg=ACCENT, relief="flat", font=("Courier", 8), cursor="hand2", padx=8, pady=4).grid(row=0, column=1)
        self.verify_rules_list = tk.Listbox(rules_lf, bg=PANEL2, fg=TEXT, font=("Courier", 7), height=3,
            relief="flat", bd=0, highlightthickness=0)
        self.verify_rules_list.grid(row=2, column=0, sticky="ew", padx=8, pady=(0, 6))
        self.verify_rules_list.bind("<Double-1>", lambda e: self._remove_verify_rule())
        self._refresh_verify_rules_list()

    def _add_verify_rule(self):
        rule = self.verify_rules_var.get().strip()
        if not rule:
            return
        rules = self.cfg.get("verify_order_rules", [])
        if rule not in rules:
            rules.append(rule)
            self.cfg["verify_order_rules"] = rules
            save_config(self.cfg)
        self.verify_rules_var.set("")
        self._refresh_verify_rules_list()

    def _remove_verify_rule(self):
        sel = self.verify_rules_list.curselection()
        if not sel:
            return
        idx = sel[0]
        rules = self.cfg.get("verify_order_rules", [])
        if 0 <= idx < len(rules):
            rules.pop(idx)
            self.cfg["verify_order_rules"] = rules
            save_config(self.cfg)
            self._refresh_verify_rules_list()

    def _refresh_verify_rules_list(self):
        self.verify_rules_list.delete(0, "end")
        for r in self.cfg.get("verify_order_rules", []):
            self.verify_rules_list.insert("end", r)

    def _run_verify_orders(self):
        pdf_path = self.verify_pdf_var.get().strip()
        excel_path = self.verify_excel_var.get().strip()

        self.verify_new_list.delete(0, "end")
        self.verify_old_list.delete(0, "end")
        self._verify_zz_descriptions = {}
        self._verify_zz_chunks = {}
        self._verify_pdf_path = ""
        self._verify_excel_path = ""

        if not pdf_path or not os.path.isfile(pdf_path):
            self.verify_summary.config(text="Please select a valid Order PDF.", fg=RED)
            return
        if not excel_path or not os.path.isfile(excel_path):
            self.verify_summary.config(text="Please select a valid Register Excel file.", fg=RED)
            return

        self.verify_summary.config(text="Loading…", fg=MUTED)
        self.update_idletasks()

        self._verify_pdf_path = pdf_path
        self._verify_excel_path = excel_path
        self._verify_customer = extract_customer_from_pdf(pdf_path) or "ASPEX"

        register_bases = get_register_zz_bases(excel_path)
        try:
            zz_with_desc = extract_zz_with_descriptions(pdf_path)
            self._verify_zz_chunks = {}
            for zz, desc, chunk in zz_with_desc:
                if zz not in self._verify_zz_descriptions:
                    self._verify_zz_descriptions[zz] = desc
                self._verify_zz_chunks[zz] = chunk
            order_zz = sorted(set(z for z, _, _ in zz_with_desc))
        except Exception as e:
            self.verify_summary.config(text=f"Error: {e}", fg=RED)
            return

        new_zz = [z for z in order_zz if z not in register_bases]
        old_zz = [z for z in order_zz if z in register_bases]
        self._verify_new_zz = list(new_zz)
        self._verify_old_zz = list(old_zz)
        self._verify_order_count = len(order_zz)
        self._verify_register_count = len(register_bases)

        for z in new_zz:
            self.verify_new_list.insert("end", z)
        for z in old_zz:
            self.verify_old_list.insert("end", z)

        if not new_zz:
            self.verify_new_list.insert("end", "(none — all parts exist)")
        if not old_zz:
            self.verify_old_list.insert("end", "(none)")

        self.verify_summary.config(
            text=f"Order: {len(order_zz)} part(s)  •  {len(new_zz)} new  •  {len(old_zz)} existing  •  Register: {len(register_bases)} part(s)",
            fg=TEXT)
        self._verify_ai_suggested_part = {}

    def _on_verify_ai_suggest(self):
        """Get AI suggestion for selected part. Stores result for use when adding."""
        sel = self.verify_new_list.curselection()
        if not sel:
            messagebox.showinfo("AI Suggest", "Select a part from the New list first.")
            return
        zz = self.verify_new_list.get(sel[0])
        if zz.startswith("("):
            return
        if not (self.cfg.get("api_key") or self.cfg.get("openai_api_key")):
            messagebox.showerror("AI Suggest", "Add an API key in Settings (Anthropic or OpenAI) first.")
            return
        zz_chunks = getattr(self, "_verify_zz_chunks", {})
        zz_desc = getattr(self, "_verify_zz_descriptions", {})
        chunk = zz_chunks.get(zz, "")
        if not chunk:
            messagebox.showwarning("AI Suggest", "No order text for this part. Run Verify Order first.")
            return
        is_door = _is_door_from_description(zz_desc.get(zz, ""))
        self.verify_summary.config(text="AI thinking…", fg=MUTED)
        self.update_idletasks()
        try:
            ai_part = ai_build_part_from_order(zz, chunk, is_door, self.cfg)
            self._verify_ai_suggested_part = {zz: ai_part} if ai_part else {}
            n_new = len(getattr(self, "_verify_new_zz", []))
            n_old = len(getattr(self, "_verify_old_zz", []))
            n_order = getattr(self, "_verify_order_count", n_new + n_old)
            n_reg = getattr(self, "_verify_register_count", 0)
            self.verify_summary.config(
                text=f"Order: {n_order} part(s)  •  {n_new} new  •  {n_old} existing  •  Register: {n_reg} part(s)",
                fg=TEXT)
            if ai_part:
                messagebox.showinfo("AI Suggest", f"Suggested:\n{ai_part}\n\nClick Add to Register to use it.")
            else:
                messagebox.showwarning("AI Suggest", "AI couldn't suggest a part code.")
        except Exception as e:
            self.verify_summary.config(text="", fg=TEXT)
            messagebox.showerror("AI Error", str(e))

    def _on_verify_add_selected(self):
        sel = self.verify_new_list.curselection()
        if not sel:
            messagebox.showinfo("Add to Register", "Select a part from the New list first.")
            return
        idx = sel[0]
        zz = self.verify_new_list.get(idx)
        if zz.startswith("("):
            return
        pdf_path = getattr(self, "_verify_pdf_path", "") or self.verify_pdf_var.get().strip()
        excel_path = getattr(self, "_verify_excel_path", "") or self.verify_excel_var.get().strip()
        zz_desc = getattr(self, "_verify_zz_descriptions", {})
        zz_chunks = getattr(self, "_verify_zz_chunks", {})

        if not pdf_path or not excel_path:
            messagebox.showerror("Add to Register", "Please run Verify Order first.")
            return

        desc = zz_desc.get(zz, "")
        is_door = _is_door_from_description(desc)
        laser_punch = "LASER" if is_door else "PUNCH"
        customer = getattr(self, "_verify_customer", None) or extract_customer_from_pdf(pdf_path) or "ASPEX"

        chunk = zz_chunks.get(zz, "")
        custom_rules = self.cfg.get("verify_order_rules", [])
        ai_suggested = getattr(self, "_verify_ai_suggested_part", {}).get(zz)
        if ai_suggested:
            part_no = ai_suggested
            self._verify_ai_suggested_part.pop(zz, None)
        else:
            part_no = build_part_description_from_order(zz, chunk, is_door, custom_rules)

        # AI fallback when regex couldn't parse (incomplete result)
        def _is_incomplete(p: str) -> bool:
            if not p or p == zz:
                return True
            if not re.search(r"\d{3,4}[xX]\d{3,4}", p):
                return True
            return False

        has_ai = bool(self.cfg.get("api_key") or self.cfg.get("openai_api_key"))
        if _is_incomplete(part_no) and has_ai and chunk:
            if messagebox.askyesno("AI Suggestion", "Couldn't parse this order fully. Use AI to suggest part description?"):
                try:
                    ai_part = ai_build_part_from_order(zz, chunk, is_door, self.cfg)
                    if ai_part:
                        part_no = ai_part
                    else:
                        messagebox.showwarning("AI", "AI couldn't suggest a part code. Using regex result.")
                except Exception as e:
                    messagebox.showerror("AI Error", str(e))

        if is_door and "_" in part_no and part_no.count("_") >= 3:
            suffix = simpledialog.askstring(
                "Add F / B / SET",
                f"Part: {part_no}\n\nEnter F (Front), B (Back), or SET:",
                parent=self,
            )
            if suffix is not None and str(suffix).strip().upper() in ("F", "B", "SET"):
                part_no = f"{part_no}_{str(suffix).strip().upper()}"

        ok, dnc = add_row_to_register_with_dnc(excel_path, laser_punch, customer, part_no, "E")
        if ok:
            self._verify_update_lists_after_add(zz)
            messagebox.showinfo("Add to Register", f"Added {part_no} as {dnc} ({laser_punch})")
        else:
            messagebox.showerror("Add to Register", "Failed to write to Excel. Ensure the file is not open elsewhere.")

    def _verify_update_lists_after_add(self, zz: str):
        """Update New/Old lists in place after adding a part — no re-scan of PDF/Excel."""
        new_zz = getattr(self, "_verify_new_zz", [])
        old_zz = getattr(self, "_verify_old_zz", [])
        if zz not in new_zz:
            return
        new_zz = [z for z in new_zz if z != zz]
        old_zz = old_zz + [zz]
        self._verify_new_zz = new_zz
        self._verify_old_zz = old_zz
        self._verify_register_count = getattr(self, "_verify_register_count", 0) + 1

        self.verify_new_list.delete(0, "end")
        self.verify_old_list.delete(0, "end")
        for z in new_zz:
            self.verify_new_list.insert("end", z)
        for z in old_zz:
            self.verify_old_list.insert("end", z)
        if not new_zz:
            self.verify_new_list.insert("end", "(none — all parts exist)")
        if not old_zz:
            self.verify_old_list.insert("end", "(none)")

        order_count = getattr(self, "_verify_order_count", len(new_zz) + len(old_zz))
        self.verify_summary.config(
            text=f"Order: {order_count} part(s)  •  {len(new_zz)} new  •  {len(old_zz)} existing  •  Register: {self._verify_register_count} part(s)",
            fg=TEXT)

    def _build_rules_banner(self):
        for w in self._rules_frame.winfo_children():
            w.destroy()
        rules = self.cfg.get("learned_rules", [])
        if not rules:
            tk.Label(self._rules_frame,
                     text="🧠  No rules loaded — add rules via chat or reset to built-ins",
                     font=("Courier", 8), fg=MUTED, bg=PANEL2,
                     padx=12, pady=7).pack(anchor="w")
        else:
            tk.Label(self._rules_frame,
                     text=f"🧠  {len(rules)} rule(s) loaded — AI uses these on every run:",
                     font=("Courier", 8, "bold"), fg=GREEN, bg=PANEL2,
                     padx=12, pady=7).pack(anchor="w")
            for r in rules[-3:]:  # show last 3
                short = r[:90] + "…" if len(r) > 90 else r
                tk.Label(self._rules_frame, text=f"  • {short}",
                         font=("Courier", 7), fg=TEXT, bg=PANEL2,
                         padx=12, pady=1, anchor="w").pack(fill="x")
            if len(rules) > 3:
                tk.Label(self._rules_frame,
                         text=f"  … and {len(rules)-3} more (ask 'Show learned rules')",
                         font=("Courier", 7), fg=MUTED, bg=PANEL2,
                         padx=12, pady=6).pack(anchor="w")
            else:
                tk.Frame(self._rules_frame, bg=PANEL2, height=6).pack()

    def _add_chat_bubble(self, role: str, text: str):
        is_user = (role == "user")

        outer = tk.Frame(self._chat_inner, bg=BG)
        outer.pack(fill="x", padx=10, pady=4)

        bubble_bg    = PANEL2 if is_user else PANEL
        bubble_fg    = TEXT
        label_text   = "You" if is_user else "AI"
        label_colour = ACCENT if is_user else PURPLE
        anchor_side  = "e" if is_user else "w"

        # label row
        lf = tk.Frame(outer, bg=BG)
        lf.pack(fill="x")
        tk.Label(lf, text=label_text,
                 font=("Courier", 7, "bold"), fg=label_colour,
                 bg=BG).pack(side="right" if is_user else "left", padx=4)

        # bubble
        bubble = tk.Frame(outer, bg=bubble_bg, padx=12, pady=8)
        bubble.pack(anchor=anchor_side, fill="x" if not is_user else None,
                    padx=(60 if is_user else 0, 0 if is_user else 60))

        msg_label = tk.Label(bubble, text=text, font=("Courier", 9),
                             fg=bubble_fg, bg=bubble_bg,
                             wraplength=480, justify="left", anchor="w")
        msg_label.pack(anchor="w", fill="x")

        # scroll to bottom
        self._chat_canvas.update_idletasks()
        self._chat_canvas.configure(scrollregion=self._chat_canvas.bbox("all"))
        self._chat_canvas.yview_moveto(1.0)

    def _add_typing_indicator(self):
        f = tk.Frame(self._chat_inner, bg=BG)
        f.pack(fill="x", padx=10, pady=4)
        tk.Label(f, text="AI", font=("Courier", 7, "bold"),
                 fg=PURPLE, bg=BG).pack(side="left", padx=4)
        inner = tk.Frame(f, bg=PANEL, padx=12, pady=8)
        inner.pack(anchor="w")
        self._typing_label = tk.Label(inner, text="● ● ●  thinking…",
                                      font=("Courier", 9), fg=MUTED, bg=PANEL)
        self._typing_label.pack()
        self._chat_canvas.update_idletasks()
        self._chat_canvas.configure(scrollregion=self._chat_canvas.bbox("all"))
        self._chat_canvas.yview_moveto(1.0)
        return f

    def _on_enter(self, event):
        if not event.state & 0x1:  # shift not held
            self._send_chat()
            return "break"

    def _quick_action(self, msg: str):
        if msg == "SYSTEM_CLEAR_RULES":
            if messagebox.askyesno("Clear rules",
                    "Clear all learned rules? The AI will start fresh."):
                self.cfg["learned_rules"] = []
                save_config(self.cfg)
                self._build_rules_banner()
                self._add_chat_bubble("assistant",
                    "✅ All learned rules cleared. Starting fresh!")
            return
        if msg == "SYSTEM_RESET_BUILTIN":
            self.cfg["learned_rules"] = list(BUILTIN_RULES)
            save_config(self.cfg)
            self._build_rules_banner()
            self._add_chat_bubble("assistant",
                f"✅ Reset to {len(BUILTIN_RULES)} built-in rules. AI knows setup sheet conventions.")
            return
        if msg == "SYSTEM_ADD_RULE":
            code = simpledialog.askstring("Add rule for unknown code",
                "Paste the part code that failed to parse:\n(e.g. ZZ9971758PXX_790X1760_S_F_LH_INTU or ZZ9971771PXX_1759X1960_D_F_INTU)")
            if not code:
                return
            desc = simpledialog.askstring("Add rule for unknown code",
                "What should the description be?\n(e.g. 790X1760 SINGLE FRAME LH INTUFRAME or 1759X1960 DOUBLE DOOR EURO INTUFRAME)")
            if not desc:
                return
            rule = f"For codes like {code.strip()}: description = {desc.strip()}"
            rules = self.cfg.get("learned_rules", [])
            rules.append(rule)
            self.cfg["learned_rules"] = rules
            save_config(self.cfg)
            self._build_rules_banner()
            self._add_chat_bubble("assistant",
                f"🧠  Rule saved: \"{rule}\"\n"
                f"I'll use this for future parsing. Run again to apply.")
            return
        self._chat_input.delete("1.0", "end")
        self._chat_input.insert("1.0", msg)
        self._send_chat()

    def _send_chat(self):
        text = self._chat_input.get("1.0", "end").strip()
        if not text: return

        self._chat_input.delete("1.0", "end")
        self._add_chat_bubble("user", text)
        self._chat_history.append({"role": "user", "content": text})

        typing_widget = self._add_typing_indicator()

        def worker():
            try:
                response = ai_chat(self._chat_history, self.cfg)

                # Check for FIX_REQUEST: in response — user asked to fix document(s)
                fix_result = None
                if "FIX_REQUEST:" in response:
                    try:
                        raw = response.split("FIX_REQUEST:")[1].split("\n")[0].strip()
                        data = json.loads(raw)
                        updates = data.get("updates") or {}
                        docs_raw = data.get("docs") or data.get("doc")
                        out_dir = getattr(self, "output_var", None)
                        out_dir = out_dir.get().strip() if out_dir else self.cfg.get("last_output", "")
                        if not out_dir:
                            out_dir = str(Path.home() / "Desktop")
                        out_dir = out_dir or "."

                        doc_refs = []
                        if docs_raw == "ALL" or (isinstance(docs_raw, str) and str(docs_raw).upper() == "ALL"):
                            for f in os.listdir(out_dir):
                                if f.lower().endswith(".docx") and not f.startswith("~"):
                                    doc_refs.append(f[:-5])  # strip .docx
                        elif isinstance(docs_raw, list):
                            doc_refs = [str(d).strip().upper().replace(".DOCX", "").replace(".PDF", "") for d in docs_raw if d]
                        elif docs_raw:
                            doc_refs = [str(docs_raw).strip().upper().replace(".DOCX", "").replace(".PDF", "")]

                        if doc_refs and updates:
                            ok_list, fail_list = [], []
                            for ref in doc_refs:
                                docx_path = os.path.join(out_dir, f"{ref}.docx")
                                if not os.path.exists(docx_path):
                                    fail_list.append(ref)
                                    continue
                                is_laser = detect_template_type(docx_path)
                                ok = update_docx_fields(docx_path, updates, is_laser, self.log)
                                if ok:
                                    pdf_path = docx_path.replace(".docx", ".pdf")
                                    if HAS_PDF:
                                        try:
                                            convert_to_pdf(docx_path, pdf_path, self.log)
                                        except Exception:
                                            pass
                                    ok_list.append(ref)
                                else:
                                    fail_list.append(ref)
                            if ok_list:
                                fix_result = f"✅ Fixed {len(ok_list)} sheet(s): " + ", ".join(ok_list)
                                if fail_list:
                                    fix_result += f"\n❌ Failed: " + ", ".join(fail_list)
                                fix_result += "\n" + ", ".join(f"{k}→{v}" for k, v in updates.items())
                            elif fail_list:
                                fix_result = f"❌ Failed: " + ", ".join(fail_list)
                        else:
                            fix_result = "❌ No documents or updates to apply"
                    except Exception as ex:
                        fix_result = f"❌ Fix error: {ex}"
                    # Use the part after FIX_REQUEST as the chat response
                    parts = response.split("FIX_REQUEST:")[1].split("\n", 1)
                    response = (parts[1].strip() if len(parts) > 1 else "") or response.split("FIX_REQUEST:")[0].strip()
                    if fix_result:
                        response = fix_result + "\n\n" + response if response else fix_result

                self._chat_history.append({"role": "assistant", "content": response})

                # Check if this looks like a correction or request to add a rule — extract and save
                correction_keywords = [
                    "wrong", "mistake", "incorrect", "should be", "not", "instead",
                    "actually", "you got", "fix", "change", "always", "never",
                    "remember", "correct", "add rule", "add a rule", "new rule",
                    "from now on", "going forward", "in future", "save this", "store this"
                ]
                is_correction = any(kw in text.lower() for kw in correction_keywords)

                rule_saved = None
                if is_correction and not fix_result:
                    rule = ai_extract_rule(text, self.cfg)
                    if rule and len(rule) > 10:
                        rules = self.cfg.get("learned_rules", [])
                        rules.append(rule)
                        self.cfg["learned_rules"] = rules
                        save_config(self.cfg)
                        rule_saved = rule

                self.after(0, lambda: self._finish_chat(
                    typing_widget, response, rule_saved))
            except Exception as e:
                self.after(0, lambda: self._finish_chat(
                    typing_widget, f"❌ Error: {e}", None))

        threading.Thread(target=worker, daemon=True).start()

    def _finish_chat(self, typing_widget, response: str, rule_saved: str | None):
        typing_widget.destroy()
        self._add_chat_bubble("assistant", response)

        if rule_saved:
            self._add_chat_bubble("assistant",
                f"🧠  Rule saved: \"{rule_saved}\"\n"
                f"I'll apply this to all future parsing and template filling.")
            self._build_rules_banner()

        # switch to generator tab if user mentioned generating/running
        trigger_words = ["generat", "run", "creat", "make the sheet"]
        if any(w in response.lower() for w in trigger_words):
            pass  # keep them on chat

    def _update_provider_badge(self):
        provider = get_provider(self.cfg)
        model    = get_model(self.cfg)
        short    = model.split("-")[0] if "-" in model else model[:8]
        if provider == "openai":
            self._provider_badge.config(text=f"OpenAI / {short}", bg=GREEN, fg=BG)
        else:
            self._provider_badge.config(text=f"Claude / {short}", bg=ACCENT, fg=BG)

    # ── browse / settings ─────────────────────────────────────────────────
    def _sync_excel_path(self, path: str):
        """Sync Register Excel path across Generator, Rename Tool, and Verify Orders tabs."""
        if path and path.strip():
            p = path.strip()
            self.excel_var.set(p)
            self.rename_excel_var.set(p)
            self.verify_excel_var.set(p)

    def _browse(self, var, filetypes):
        path = filedialog.askopenfilename(filetypes=filetypes) if filetypes \
               else filedialog.askdirectory()
        if path:
            var.set(path)
            if var is self.excel_var:
                self._sync_excel_path(path)

    def _show_about(self):
        messagebox.showinfo(
            "About SheetGen",
            "SheetGen\n\n"
            "Copyright © 2026 Pruthvi. All rights reserved.\n\n"
            "Unauthorized copying, distribution, or modification of this software is prohibited.",
        )

    def _open_settings(self):
        win = tk.Toplevel(self)
        win.title("Settings")
        win.geometry("520x480")
        win.configure(bg=BG)
        win.grab_set()
        win.resizable(False, False)

        tk.Label(win, text="DEFAULTS",
                 font=("Courier", 11, "bold"), fg=TEXT, bg=BG).pack(
                     anchor="w", padx=20, pady=(18, 2))
        tk.Frame(win, bg=ACCENT, height=1).pack(fill="x", padx=20, pady=(0, 8))
        prog_frame = tk.Frame(win, bg=BG)
        prog_frame.pack(fill="x", padx=20, pady=(0, 8))
        prog_frame.columnconfigure(1, weight=1)
        tk.Label(prog_frame, text="Programmer", font=("Courier", 9), fg=MUTED, bg=BG, width=14).pack(side="left")
        prog_var = tk.StringVar(value=self.cfg.get("programmer", ""))
        tk.Entry(prog_frame, textvariable=prog_var, bg=PANEL2, fg=TEXT, font=("Courier", 9),
                 relief="flat", insertbackground=ACCENT).pack(side="left", fill="x", expand=True, padx=(0,10), ipady=5)

        tk.Label(win, text="AI PROVIDER SETTINGS",
                 font=("Courier", 11, "bold"), fg=TEXT, bg=BG).pack(
                     anchor="w", padx=20, pady=(18, 2))
        tk.Frame(win, bg=ACCENT, height=1).pack(fill="x", padx=20, pady=(0, 14))

        # ── provider selector ──
        prov_frame = tk.Frame(win, bg=BG)
        prov_frame.pack(fill="x", padx=20, pady=(0, 12))
        tk.Label(prov_frame, text="Provider", font=("Courier", 9, "bold"),
                 fg=MUTED, bg=BG, width=14, anchor="w").pack(side="left")
        prov_var = tk.StringVar(value=self.cfg.get("ai_provider", "anthropic"))
        for val, label in [("anthropic", "Anthropic (Claude)"),
                           ("openai",    "OpenAI (GPT)")]:
            tk.Radiobutton(prov_frame, text=label, variable=prov_var, value=val,
                           bg=BG, fg=TEXT, selectcolor=PANEL2,
                           activebackground=BG, activeforeground=ACCENT,
                           font=("Courier", 9), cursor="hand2").pack(side="left", padx=(0,16))

        tk.Frame(win, bg=BORDER, height=1).pack(fill="x", padx=20, pady=8)

        # ── Anthropic section ──
        anth_frame = tk.LabelFrame(win, text=" Anthropic ",
                                   font=("Courier", 8, "bold"),
                                   fg=ACCENT, bg=BG, bd=1,
                                   highlightbackground=BORDER)
        anth_frame.pack(fill="x", padx=20, pady=(0, 10))
        anth_frame.columnconfigure(1, weight=1)

        tk.Label(anth_frame, text="API Key", font=("Courier", 8),
                 fg=MUTED, bg=BG, width=10, anchor="w").grid(
                     row=0, column=0, sticky="w", padx=10, pady=(8,4))
        anth_key_var = tk.StringVar(value=self.cfg.get("api_key", ""))
        tk.Entry(anth_frame, textvariable=anth_key_var, show="•",
                 bg=PANEL2, fg=TEXT, font=("Courier", 9), relief="flat",
                 insertbackground=ACCENT, bd=0,
                 highlightthickness=1, highlightbackground=BORDER,
                 highlightcolor=ACCENT).grid(
                     row=0, column=1, sticky="ew", padx=(0,10), pady=(8,4), ipady=5)

        tk.Label(anth_frame, text="Model", font=("Courier", 8),
                 fg=MUTED, bg=BG, width=10, anchor="w").grid(
                     row=1, column=0, sticky="w", padx=10, pady=(0,8))
        anth_model_var = tk.StringVar(value=self.cfg.get("anthropic_model", ANTHROPIC_MODELS[0]))
        anth_model_dd  = ttk.Combobox(anth_frame, textvariable=anth_model_var,
                                      values=ANTHROPIC_MODELS, state="readonly",
                                      font=("Courier", 8))
        anth_model_dd.grid(row=1, column=1, sticky="ew", padx=(0,10), pady=(0,8), ipady=3)

        tk.Label(anth_frame,
                 text="Get key → console.anthropic.com",
                 font=("Courier", 7), fg=MUTED, bg=BG).grid(
                     row=2, column=0, columnspan=2, sticky="w", padx=10, pady=(0,6))

        # ── OpenAI section ──
        oai_frame = tk.LabelFrame(win, text=" OpenAI ",
                                  font=("Courier", 8, "bold"),
                                  fg=GREEN, bg=BG, bd=1,
                                  highlightbackground=BORDER)
        oai_frame.pack(fill="x", padx=20, pady=(0, 10))
        oai_frame.columnconfigure(1, weight=1)

        tk.Label(oai_frame, text="API Key", font=("Courier", 8),
                 fg=MUTED, bg=BG, width=10, anchor="w").grid(
                     row=0, column=0, sticky="w", padx=10, pady=(8,4))
        oai_key_var = tk.StringVar(value=self.cfg.get("openai_api_key", ""))
        tk.Entry(oai_frame, textvariable=oai_key_var, show="•",
                 bg=PANEL2, fg=TEXT, font=("Courier", 9), relief="flat",
                 insertbackground=GREEN, bd=0,
                 highlightthickness=1, highlightbackground=BORDER,
                 highlightcolor=GREEN).grid(
                     row=0, column=1, sticky="ew", padx=(0,10), pady=(8,4), ipady=5)

        tk.Label(oai_frame, text="Model", font=("Courier", 8),
                 fg=MUTED, bg=BG, width=10, anchor="w").grid(
                     row=1, column=0, sticky="w", padx=10, pady=(0,8))
        oai_model_var = tk.StringVar(value=self.cfg.get("openai_model", OPENAI_MODELS[0]))
        oai_model_dd  = ttk.Combobox(oai_frame, textvariable=oai_model_var,
                                     values=OPENAI_MODELS, state="readonly",
                                     font=("Courier", 8))
        oai_model_dd.grid(row=1, column=1, sticky="ew", padx=(0,10), pady=(0,8), ipady=3)

        tk.Label(oai_frame,
                 text="Get key → platform.openai.com",
                 font=("Courier", 7), fg=MUTED, bg=BG).grid(
                     row=2, column=0, columnspan=2, sticky="w", padx=10, pady=(0,6))

        # ── save ──
        def _save():
            prog = prog_var.get().strip()
            self.cfg["programmer"] = prog
            if hasattr(self, "programmer_var"):
                self.programmer_var.set(prog)
            self.cfg["ai_provider"]     = prov_var.get()
            self.cfg["api_key"]         = anth_key_var.get().strip()
            self.cfg["anthropic_model"] = anth_model_var.get()
            self.cfg["openai_api_key"]  = oai_key_var.get().strip()
            self.cfg["openai_model"]    = oai_model_var.get()
            save_config(self.cfg)
            win.destroy()
            provider_label = "Anthropic" if self.cfg["ai_provider"] == "anthropic" else "OpenAI"
            model_label    = self.cfg["anthropic_model"] if self.cfg["ai_provider"] == "anthropic" \
                             else self.cfg["openai_model"]
            self.log(f"✅ Settings saved — using {provider_label} / {model_label}")
            self._update_provider_badge()

        tk.Button(win, text="💾  Save Settings", command=_save,
                  bg=ACCENT, fg=BG, font=("Courier", 10, "bold"),
                  relief="flat", padx=20, pady=9, cursor="hand2").pack(pady=(4, 16))

    # ── log ───────────────────────────────────────────────────────────────
    def log(self, msg: str):
        self.log_box.configure(state="normal")
        tag = ("err"  if any(x in msg for x in ["❌","Error","error","ERROR"]) else
               "ok"   if any(x in msg for x in ["✅","🎉","saved","PDF","💾"]) else
               "warn" if any(x in msg for x in ["⚠️","warn"]) else "")
        self.log_box.insert("end", msg + "\n", tag)
        self.log_box.see("end"); self.log_box.configure(state="disabled")
        self.update_idletasks()

    # ── generate ──────────────────────────────────────────────────────────
    def _run(self):
        ep = self.excel_var.get().strip()
        lp = self.laser_tpl_var.get().strip()
        cp = self.combi_tpl_var.get().strip()
        od = self.output_var.get().strip()
        if not ep or not os.path.exists(ep):
            messagebox.showerror("Missing","Select a valid Excel file."); return
        if not lp or not os.path.exists(lp):
            messagebox.showerror("Missing","Select a valid Laser template."); return
        if not cp or not os.path.exists(cp):
            messagebox.showerror("Missing","Select a valid Combi template."); return
        if not od:
            messagebox.showerror("Missing","Select an output folder."); return

        start_row = max(1, self.start_row_var.get())
        max_rows = max(1, self.max_rows_var.get())

        self.cfg.update({
            "last_excel": ep, "last_laser_template": lp, "last_combi_template": cp,
            "last_output": od, "start_row": start_row, "max_rows": max_rows,
            "programmer": self.programmer_var.get(),
            "sheet_size": self.sheet_size_var.get(), "pts_prog": self.pts_prog_var.get(),
            "pts_sht": self.pts_sht_var.get(), "thickness": self.thickness_var.get(),
            "material_type": self.material_type_var.get(),
            "tape": self.tape_var.get(), "inspect": self.inspect_var.get(),
            "su_sheet": self.su_sheet_var.get(),
            "use_ai_fallback": self.use_ai_fallback_var.get(),
        })
        # Update active profile with current values
        ap = self.cfg.get("active_profile")
        if ap and ap in self.cfg.get("profiles", {}):
            self.cfg["profiles"][ap] = self._get_current_ui_data()
        save_config(self.cfg)

        self.gen_btn.config(state="disabled")
        self.progress.start(10)
        self.log("\n" + "─"*50)
        self.log("🚀 Starting generation…")
        threading.Thread(target=self._worker,
                         args=(ep, lp, cp, od, start_row, max_rows), daemon=True).start()

    def _worker(self, ep, lp, cp, od, start_row, max_rows):
        try:
            # Prefer selected search results; else use filter mode
            rows = self._get_selected_search_rows()
            if rows:
                self.log(f"📋 Using {len(rows)} selected search result(s)")
            else:
                self.log("📊 No search selection — using filter mode (Rows, Filter, Max scan)…")
                fv = (self.filter_var.get() or "All").strip()
                filter_type = "" if fv == "All" else ("LASER" if "LASER" in fv else "PUNCH")
                max_scan = self._safe_max_scan()
                rows = read_register_excel(ep, start_row, max_rows, self.log,
                                           filter_type=filter_type, max_scan=max_scan)
            if not rows:
                self.log("❌ No data rows found."); return

            opts = {
                "programmer": self.programmer_var.get(),
                "sheet_size": self.sheet_size_var.get(),
                "pts_prog": self.pts_prog_var.get(),
                "pts_sht": self.pts_sht_var.get(),
                "thickness": self.thickness_var.get(),
                "material_type": self.material_type_var.get(),
                "tape": self.tape_var.get() or "Choose an item.",
                "inspect": self.inspect_var.get() or "Choose an item.",
                "su_sheet": self.su_sheet_var.get() or "Choose an item.",
                "laser_ticks": [v.get() for v in self.laser_ticks],
                "combi_ticks": [v.get() for v in getattr(self, "combi_ticks", [])],
            }

            os.makedirs(od, exist_ok=True)
            success = 0
            for idx, row in enumerate(rows):
                part_code = row.get("part_code", "")
                dnc = row.get("dnc", "") or f"JOB_{idx+1:03d}"
                self.log(f"\n── Row {idx+1}: {part_code}")

                parsed = parse_zz_part_code(part_code)
                parsed.update(row)
                parsed["revision"] = row.get("revision", "E")
                parsed["assy_number"] = ""  # blank for ZZ orders

                # AI fallback when rules struggle (empty description, missing part_number, etc.)
                use_ai = getattr(self, "use_ai_fallback_var", None)
                has_key = bool(self.cfg.get("api_key") or self.cfg.get("openai_api_key"))
                if use_ai and use_ai.get() and has_key:
                    needs_ai = (
                        not parsed.get("description") or parsed.get("description") == part_code
                        or not parsed.get("part_number") or not parsed.get("dimensions")
                    )
                    if needs_ai:
                        try:
                            ai_parsed = ai_parse_code(part_code, self.cfg, self.log)
                            for k, v in ai_parsed.items():
                                if v and (not parsed.get(k) or parsed.get(k) == part_code):
                                    parsed[k] = v
                            if not parsed.get("description") or parsed.get("description") == part_code:
                                parsed["description"] = ai_parsed.get("description", parsed.get("description", part_code))
                        except Exception as e:
                            self.log(f"  ⚠️ AI fallback skipped: {e}")

                # Hardcoded: LASER → laser sheet, PUNCH/COMBI → combi sheet (no AI)
                is_laser, tp = pick_template(row.get("laser_punch"), lp, cp)
                opts["is_laser"] = is_laser
                self.log(f"  📄 {'Laser' if is_laser else 'Combi (PUNCH)'} template")

                safe = re.sub(r'[^\w\-_.]', '_', str(dnc))
                docx_out = os.path.join(od, f"{safe}.docx")
                pdf_out = os.path.join(od, f"{safe}.pdf")

                try:
                    fill_word_template(tp, docx_out, parsed, opts, self.log)
                except Exception as e:
                    self.log(f"  ❌ Fill error: {e}"); continue

                if self.save_pdf_var.get():
                    try: convert_to_pdf(docx_out, pdf_out, self.log)
                    except Exception as e: self.log(f"  ⚠️  PDF: {e}")

                if not self.save_docx_var.get():
                    try: os.remove(docx_out)
                    except: pass

                success += 1

            self.log(f"\n🎉 Done! {success}/{len(rows)} sheets → {od}")

        except Exception as e:
            self.log(f"\n❌ Fatal: {e}")
            import traceback; self.log(traceback.format_exc())
        finally:
            self.after(0, self._done)

    def _done(self):
        self.progress.stop()
        self.gen_btn.config(state="normal")


if __name__ == "__main__":
    app = App()
    app.mainloop()
